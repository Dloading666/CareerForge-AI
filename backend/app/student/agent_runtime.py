from __future__ import annotations

import base64
import json
import logging
import time
import mimetypes
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, AsyncIterator, Optional

import httpx
from fastapi import HTTPException, UploadFile, status
import re as _re

from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.admin.master_service import DEFAULT_SYSTEM_PROMPT, get_or_create_master_config
from app.admin.model_service import decrypt_api_key
from app.admin.models import ModelConfig
from app.auth.models import StudentUser
from app.auth.service import AuthIdentity
from app.core.config import get_settings
from app.skills.service import list_skills, serialize_skill
from app.student.agent_models import (
    StudentAgentActivity,
    StudentAgentAttachment,
    StudentAgentMessage,
    StudentAgentSession,
)
from app.student.agent_schemas import AgentActivityResponse, AgentAttachmentResponse, AgentModelOptionResponse
from app.student.resume_models import StudentResume
from app.student.tool_validation import parse_tool_arguments

logger = logging.getLogger(__name__)


# ── Value objects ──────────────────────────────────────────────────────────────


@dataclass
class RuntimeObservation:
    kind: str
    name: str
    summary: str
    detail: dict[str, Any]


@dataclass
class ToolDefinition:
    name: str
    description: str
    source: str
    priority: int
    input_schema: dict[str, Any]
    metadata: dict[str, Any]


# ── Constants ──────────────────────────────────────────────────────────────────

# Capabilities that can serve OpenAI-compatible chat completions for the master
# agent. The 模型广场 only tags models as "text" / "multimodal" (there is no
# "chat" option in the admin form), so the student side must accept those — plus
# "chat" for backward compatibility. Embedding / rerank models are excluded.
CHAT_CAPABLE_CAPABILITIES = ("text", "multimodal", "chat")

# TTS 模型仅供面试官类智能体使用，与聊天模型互斥
TTS_CAPABLE_CAPABILITIES = ("tts",)
INTERVIEW_AGENT_CATEGORIES = ("interview",)


def _agent_allowed_capabilities(category: str | None) -> tuple[str, ...]:
    """根据智能体类别返回允许使用的模型 capability 集合。"""
    if category in INTERVIEW_AGENT_CATEGORIES:
        return TTS_CAPABLE_CAPABILITIES
    return CHAT_CAPABLE_CAPABILITIES
AUTO_ATTACHMENT_PROMPT = "请帮我分析上传的附件。"


# ── Utilities ──────────────────────────────────────────────────────────────────


def utcnow() -> datetime:
    return datetime.now(timezone.utc)


def dumps_event(event: str, data: dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


# ── Built-in tool definitions ──────────────────────────────────────────────────


BUILTIN_TOOLS: list[ToolDefinition] = [
    ToolDefinition(
        name="invoke_agent",
        description="调用主智能体工具注册表中的子智能体，可派发给面试官、简历优化、岗位匹配等团队成员。",
        source="builtin",
        priority=1000,
        input_schema={
            "type": "object",
            "properties": {"agent_key": {"type": "string"}, "task": {"type": "string"}},
            "required": ["agent_key", "task"],
        },
        metadata={"kind": "subagent"},
    ),
    ToolDefinition(
        name="query_student_profile",
        description="查询学生档案，包括姓名、邮箱、学院、专业、年级等基础背景。",
        source="builtin",
        priority=990,
        input_schema={"type": "object", "properties": {}, "required": []},
        metadata={"kind": "profile"},
    ),
    ToolDefinition(
        name="query_job_positions",
        description="搜索岗位库，用于岗位匹配、JD 查询和职位推荐。",
        source="builtin",
        priority=980,
        input_schema={
            "type": "object",
            "properties": {"keyword": {"type": "string"}, "company": {"type": "string"}, "role": {"type": "string"}},
            "required": [],
        },
        metadata={"kind": "job"},
    ),
    ToolDefinition(
        name="query_knowledge_base",
        description="检索就业政策、行业知识、公司简介等知识库内容。",
        source="builtin",
        priority=970,
        input_schema={"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]},
        metadata={"kind": "knowledge"},
    ),
    ToolDefinition(
        name="read_resume",
        description="读取或解析学生上传的简历文件。当前版本在未上传文件时返回材料缺口。",
        source="builtin",
        priority=960,
        input_schema={"type": "object", "properties": {"file_id": {"type": "string"}}, "required": []},
        metadata={"kind": "resume"},
    ),
    ToolDefinition(
        name="analyze_uploaded_file",
        description="分析学生上传的图片、Word、PDF、Excel、文本等附件，并把提取内容交给主智能体综合。",
        source="builtin",
        priority=955,
        input_schema={"type": "object", "properties": {"attachment_ids": {"type": "array"}}, "required": []},
        metadata={"kind": "file"},
    ),
    ToolDefinition(
        name="send_notification",
        description="发送邮件或站内通知，用于面试提醒、报告推送等需要学生确认的动作。",
        source="builtin",
        priority=950,
        input_schema={
            "type": "object",
            "properties": {"channel": {"type": "string"}, "content": {"type": "string"}},
            "required": ["content"],
        },
        metadata={"kind": "notification", "risk": "medium"},
    ),
    ToolDefinition(
        name="get_session_context",
        description="读取当前会话历史，用于 ReAct Observe 阶段回溯上下文。",
        source="builtin",
        priority=940,
        input_schema={"type": "object", "properties": {"limit": {"type": "integer"}}, "required": []},
        metadata={"kind": "context"},
    ),
    ToolDefinition(
        name="export_resume_pdf",
        description=(
            "把优化后的简历内容生成一份可下载的 PDF 文件，并返回下载链接。"
            "仅在你已经基于学生的真实简历（通过 read_resume 读取）完成改写后调用；"
            "禁止凭空捏造经历来生成简历。"
        ),
        source="builtin",
        priority=965,
        input_schema={
            "type": "object",
            "properties": {
                "markdown": {"type": "string", "description": "完整的简历正文，Markdown 格式（标题用 #，要点用 -）。"},
                "filename": {"type": "string", "description": "文件名，可选，例如『张三-后端简历』。"},
            },
            "required": ["markdown"],
        },
        metadata={"kind": "resume", "risk": "low"},
    ),
    ToolDefinition(
        name="read_webpage",
        description="读取指定 URL 的网页内容，返回 Markdown 格式的正文。适用于学生发送链接、需要查看招聘信息、公司官网等场景。",
        source="builtin",
        priority=900,
        input_schema={
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "要读取的网页 URL"},
                "max_length": {"type": "integer", "description": "返回内容最大字符数，默认 5000"},
            },
            "required": ["url"],
        },
        metadata={"kind": "web"},
    ),
    ToolDefinition(
        name="web_search",
        description="联网搜索关键词，返回搜索结果摘要。适用于查询公司背景、行业动态、岗位信息等需要实时网络数据的场景。",
        source="builtin",
        priority=895,
        input_schema={
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "搜索关键词"},
                "num_results": {"type": "integer", "description": "返回结果数量，默认 5"},
            },
            "required": ["query"],
        },
        metadata={"kind": "web"},
    ),
    ToolDefinition(
        name="generate_resume_data",
        description=(
            "根据学生信息和目标 JD，生成一份结构化在线简历并保存到系统。"
            "调用前必须先 query_student_profile 读取学生信息。"
            "调用成功后会返回 editor_url，用 Markdown 链接格式呈现给学生。"
        ),
        source="builtin",
        priority=970,
        input_schema={
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "简历标题，例如『张三-后端工程师简历』"},
                "template_id": {"type": "string", "description": "模板ID: classic/modern/elegant，默认 classic"},
                "basic": {
                    "type": "object",
                    "description": "基本信息",
                    "properties": {
                        "name": {"type": "string"},
                        "target_position": {"type": "string", "description": "目标职位"},
                        "email": {"type": "string"},
                        "phone": {"type": "string"},
                        "location": {"type": "string"},
                        "birth_date": {"type": "string", "description": "格式 YYYY-MM"},
                    },
                },
                "education": {
                    "type": "array",
                    "description": "教育经历列表",
                    "items": {
                        "type": "object",
                        "properties": {
                            "school": {"type": "string"},
                            "major": {"type": "string"},
                            "degree": {"type": "string"},
                            "start_date": {"type": "string"},
                            "end_date": {"type": "string"},
                            "gpa": {"type": "string"},
                            "description": {"type": "string", "description": "每行一个亮点，换行分隔"},
                        },
                    },
                },
                "experience": {
                    "type": "array",
                    "description": "工作经历列表",
                    "items": {
                        "type": "object",
                        "properties": {
                            "company": {"type": "string"},
                            "position": {"type": "string"},
                            "date": {"type": "string", "description": "时间段，例如 2022.06 - 2024.12"},
                            "details": {"type": "string", "description": "每行一个要点，换行分隔"},
                        },
                    },
                },
                "projects": {
                    "type": "array",
                    "description": "项目经历列表",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "role": {"type": "string"},
                            "date": {"type": "string"},
                            "description": {"type": "string", "description": "每行一个要点，换行分隔"},
                        },
                    },
                },
                "skills": {"type": "string", "description": "技能描述，每行一条，换行分隔"},
                "self_evaluation": {"type": "string", "description": "自我评价，每行一段，换行分隔"},
            },
            "required": ["title", "basic"],
        },
        metadata={"kind": "resume"},
    ),
    ToolDefinition(
        name="optimize_resume_data",
        description=(
            "基于学生已有简历内容和目标 JD，生成一份优化版简历并保存到系统。"
            "调用前必须先 read_resume 读取学生简历内容，禁止凭空捏造。"
            "调用成功后会返回 editor_url，用 Markdown 链接格式呈现给学生。"
        ),
        source="builtin",
        priority=968,
        input_schema={
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "优化后简历标题"},
                "template_id": {"type": "string", "description": "模板ID: classic/modern/elegant"},
                "source_resume_id": {"type": "integer", "description": "来源的在线简历 ID（如有）"},
                "basic": {"type": "object"},
                "education": {"type": "array", "items": {"type": "object"}},
                "experience": {"type": "array", "items": {"type": "object"}},
                "projects": {"type": "array", "items": {"type": "object"}},
                "skills": {"type": "string"},
                "self_evaluation": {"type": "string"},
            },
            "required": ["title", "basic"],
        },
        metadata={"kind": "resume"},
    ),
    ToolDefinition(
        name="update_resume_data",
        description=(
            "更新学生已有的在线简历（局部修改）。"
            "调用前必须先 read_resume 确认简历内容，需要 resume_id。"
        ),
        source="builtin",
        priority=966,
        input_schema={
            "type": "object",
            "properties": {
                "resume_id": {"type": "integer", "description": "要更新的简历 ID"},
                "title": {"type": "string"},
                "template_id": {"type": "string"},
                "basic": {"type": "object"},
                "education": {"type": "array", "items": {"type": "object"}},
                "experience": {"type": "array", "items": {"type": "object"}},
                "projects": {"type": "array", "items": {"type": "object"}},
                "skills": {"type": "string"},
                "self_evaluation": {"type": "string"},
            },
            "required": ["resume_id"],
        },
        metadata={"kind": "resume"},
    ),
]


def _tool_safe_name(value: str) -> str:
    clean = "".join(ch if ch.isalnum() else "_" for ch in value.lower()).strip("_")
    return clean or "skill_tool"


# ── Attachment handling ────────────────────────────────────────────────────────


def _is_allowed_attachment(ext: str, content_type: str) -> bool:
    allowed_ext = {
        "png", "jpg", "jpeg", "webp", "gif",
        "pdf", "docx", "doc", "xlsx", "xls",
        "csv", "txt", "md", "json",
    }
    if ext in allowed_ext:
        return True
    return content_type.startswith("image/")


def _extract_attachment_text(path: Path, content_type: str, ext: str) -> str:
    try:
        if ext == "pdf":
            return _extract_pdf_text(path)
        if ext == "docx":
            return _extract_docx_text(path)
        if ext in {"xlsx", "xls"}:
            return _extract_xlsx_text(path)
        if ext in {"csv", "txt", "md", "json"}:
            return path.read_text(encoding="utf-8", errors="ignore")[:12000]
        if content_type.startswith("image/"):
            return _extract_image_summary(path)
    except Exception as exc:
        logger.exception("附件解析失败: %s", path)
        return f"附件已保存，但自动解析失败：{str(exc)[:200]}"
    return "附件已保存，当前格式需要专用 Skill 或外部工具进一步解析。"


def _extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for index, page in enumerate(reader.pages[:12], start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append(f"[PDF 第 {index} 页]\n{text}")
    return "\n\n".join(chunks)[:12000] or "PDF 未提取到可读文本，可能是扫描件。"


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    chunks = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables[:8]:
        for row in table.rows[:30]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks)[:12000] or "Word 文档未提取到可读文本。"


def _extract_xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    chunks: list[str] = []
    for sheet in workbook.worksheets[:5]:
        chunks.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(min_row=1, max_row=40, max_col=12, values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks)[:12000] or "Excel 文件未提取到可读内容。"


def _extract_image_summary(path: Path) -> str:
    from PIL import Image

    with Image.open(path) as image:
        width, height = image.size
        mode = image.mode
    return f"图片附件已保存：{width}x{height}，色彩模式 {mode}。如所选模型支持视觉输入，将随请求一并传入。"


# ── Session CRUD ───────────────────────────────────────────────────────────────


def create_session(db: Session, identity: AuthIdentity, title: Optional[str]) -> StudentAgentSession:
    session = StudentAgentSession(
        tenant_id=identity.tenant_id,
        student_id=identity.user_id,
        title=(title or "新对话").strip() or "新对话",
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return session


def list_available_models(
    db: Session,
    identity: AuthIdentity,
    allowed_capabilities: tuple[str, ...] = CHAT_CAPABLE_CAPABILITIES,
) -> list[AgentModelOptionResponse]:
    rows = db.scalars(
        select(ModelConfig)
        .where(
            ModelConfig.tenant_id == identity.tenant_id,
            ModelConfig.is_deleted.is_(False),
            ModelConfig.open_to_student.is_(True),
            ModelConfig.capability.in_(allowed_capabilities),
            ModelConfig.status == "active",
        )
        .order_by(ModelConfig.id.asc())
    ).all()
    return [AgentModelOptionResponse.model_validate(row) for row in rows]


def serialize_attachment(attachment: StudentAgentAttachment) -> AgentAttachmentResponse:
    data = AgentAttachmentResponse.model_validate(attachment)
    data.download_url = _attachment_download_url(attachment.stored_path)
    return data


async def save_attachment(
    db: Session,
    identity: AuthIdentity,
    session_id: int,
    upload: UploadFile,
) -> StudentAgentAttachment:
    session = get_session_or_404(db, identity, session_id)
    settings = get_settings()
    raw = await upload.read()
    if not raw:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="附件为空")
    if len(raw) > settings.agent_upload_max_bytes:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="附件超过大小限制")

    original_name = Path(upload.filename or "attachment").name
    ext = Path(original_name).suffix.lower().lstrip(".")
    content_type = upload.content_type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"
    if not _is_allowed_attachment(ext, content_type):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="暂不支持该附件格式")

    storage_dir = Path(settings.agent_upload_storage_dir) / str(identity.tenant_id) / str(identity.user_id)
    storage_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid.uuid4().hex}.{ext or 'bin'}"
    stored_path = storage_dir / stored_name
    stored_path.write_bytes(raw)

    extracted_text = _extract_attachment_text(stored_path, content_type, ext)
    row = StudentAgentAttachment(
        tenant_id=identity.tenant_id,
        student_id=identity.user_id,
        session_id=session.id,
        original_name=original_name,
        stored_path=str(stored_path),
        content_type=content_type,
        file_ext=ext,
        file_size=len(raw),
        extracted_text=extracted_text,
        status="ready",
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return row


def list_sessions(db: Session, identity: AuthIdentity) -> list[StudentAgentSession]:
    # 只返回「至少有一条消息」的会话，自动隐藏从未对话过的空会话
    has_message = (
        select(StudentAgentMessage.id)
        .where(StudentAgentMessage.session_id == StudentAgentSession.id)
        .exists()
    )
    return list(
        db.scalars(
            select(StudentAgentSession)
            .where(
                StudentAgentSession.tenant_id == identity.tenant_id,
                StudentAgentSession.student_id == identity.user_id,
                StudentAgentSession.status == "active",
                has_message,
            )
            .order_by(StudentAgentSession.updated_at.desc())
        ).all()
    )


def delete_session(db: Session, identity: AuthIdentity, session_id: int) -> None:
    session = get_session_or_404(db, identity, session_id)
    session.status = "deleted"
    session.updated_at = utcnow()
    db.commit()


def get_session_or_404(db: Session, identity: AuthIdentity, session_id: int) -> StudentAgentSession:
    session = db.scalar(
        select(StudentAgentSession).where(
            StudentAgentSession.id == session_id,
            StudentAgentSession.tenant_id == identity.tenant_id,
            StudentAgentSession.student_id == identity.user_id,
            StudentAgentSession.status == "active",
        )
    )
    if session is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="对话不存在")
    return session


def get_history(db: Session, identity: AuthIdentity, session_id: int):
    session = get_session_or_404(db, identity, session_id)
    messages = list(
        db.scalars(
            select(StudentAgentMessage)
            .where(StudentAgentMessage.session_id == session.id)
            .order_by(StudentAgentMessage.id.asc())
        ).all()
    )
    activities = list(
        db.scalars(
            select(StudentAgentActivity)
            .where(StudentAgentActivity.session_id == session.id)
            .order_by(StudentAgentActivity.id.asc())
        ).all()
    )
    attachments = list(
        db.scalars(
            select(StudentAgentAttachment)
            .where(StudentAgentAttachment.session_id == session.id)
            .order_by(StudentAgentAttachment.id.asc())
        ).all()
    )
    return session, messages, activities, attachments


# ── DB helpers ─────────────────────────────────────────────────────────────────


def serialize_activity(activity: StudentAgentActivity) -> AgentActivityResponse:
    detail: dict[str, Any] = {}
    if activity.detail_json:
        try:
            detail = json.loads(activity.detail_json)
        except json.JSONDecodeError:
            detail = {}
    return AgentActivityResponse(
        id=activity.id,
        session_id=activity.session_id,
        message_id=activity.message_id,
        kind=activity.kind,
        name=activity.name,
        status=activity.status,
        summary=activity.summary,
        detail=detail,
        started_at=activity.started_at,
        completed_at=activity.completed_at,
    )


def _save_message(db: Session, session: StudentAgentSession, role: str, content: str) -> StudentAgentMessage:
    message = StudentAgentMessage(session_id=session.id, role=role, content=content)
    session.updated_at = utcnow()
    if role == "user" and session.title == "新对话":
        session.title = content.strip().replace("\n", " ")[:32] or "新对话"
    db.add(message)
    db.commit()
    db.refresh(message)
    return message


def _save_activity(
    db: Session,
    session: StudentAgentSession,
    message: StudentAgentMessage,
    *,
    kind: str,
    name: str,
    status_value: str,
    summary: str,
    detail: Optional[dict[str, Any]] = None,
) -> StudentAgentActivity:
    activity = StudentAgentActivity(
        session_id=session.id,
        message_id=message.id,
        kind=kind,
        name=name,
        status=status_value,
        summary=summary,
        detail_json=json.dumps(detail or {}, ensure_ascii=False),
        completed_at=utcnow() if status_value in {"completed", "failed"} else None,
    )
    db.add(activity)
    db.commit()
    db.refresh(activity)
    return activity


def _complete_activity(
    db: Session,
    activity: StudentAgentActivity,
    *,
    status_value: str,
    summary: str,
    detail: dict[str, Any],
) -> StudentAgentActivity:
    activity.status = status_value
    activity.summary = summary
    activity.detail_json = json.dumps(detail, ensure_ascii=False)
    activity.completed_at = utcnow()
    db.commit()
    db.refresh(activity)
    return activity


# ── Main streaming entry point ─────────────────────────────────────────────────


async def stream_master_reply(
    db: Session,
    identity: AuthIdentity,
    session_id: int,
    content: str,
    model_id: Optional[int],
    reasoning_effort: str,
    attachment_ids: list[int],
) -> AsyncIterator[str]:
    """Agentic-loop entry point.

    Harness owns the loop: the model only proposes tool calls, the Harness
    validates / executes / audits them and feeds results back, until the model
    produces a final answer or `max_iterations` is reached. See the in-repo
    《Agent = Model + Harness 开发准则》— "Harness 提供信任".
    """
    session = get_session_or_404(db, identity, session_id)
    user_message = _save_message(db, session, "user", content.strip())
    attachments = _claim_message_attachments(db, identity, session, user_message, attachment_ids)
    if (
        content.strip() == AUTO_ATTACHMENT_PROMPT
        and attachments
        and all(attachment.content_type.startswith("image/") for attachment in attachments)
        and session.title == AUTO_ATTACHMENT_PROMPT
    ):
        session.title = "图片分析"
        db.commit()
    yield dumps_event("message.saved", {"message_id": user_message.id})

    model = _select_chat_model(db, identity.tenant_id, model_id)

    # ── Model availability guards (return a controlled assistant message) ──
    if model is None or not model.api_key_cipher:
        assistant_message = StudentAgentMessage(session_id=session.id, role="assistant", content="")
        db.add(assistant_message)
        db.commit()
        db.refresh(assistant_message)
        if model is None:
            error = "当前没有可用的聊天模型，请管理员在模型广场开启「对学生开放」。"
        else:
            error = f"模型「{model.display_name}」未配置 API Key，请管理员在模型广场补全配置。"
        assistant_message.content = error
        session.updated_at = utcnow()
        db.commit()
        yield dumps_event("message.delta", {"message_id": assistant_message.id, "delta": error})
        yield dumps_event("message.completed", {"message_id": assistant_message.id})
        yield dumps_event("done", {"session_id": session.id})
        return

    config = get_or_create_master_config(db, identity.tenant_id)
    # Harness hard boundary — 尊重管理端配置的轮次，但保留一个安全上限防止失控。
    max_iterations = max(1, min(int(config.max_iterations or 8), 20))
    permission_mode = (config.permission_mode or "ask").lower()

    # Curated, safe tool registry. Only tools the Harness can honestly fulfil
    # are exposed — fabricating stubs are intentionally excluded.
    tool_defs = assemble_active_tools(db, identity)
    registry = {tool.name: tool for tool in tool_defs}
    openai_tools = _build_openai_tools(tool_defs)

    # Build initial messages BEFORE creating the empty assistant row, so the
    # history loader does not pick up a blank assistant turn.
    messages = _build_initial_messages(
        db, identity, session, content, reasoning_effort, model, attachments, config
    )

    assistant_message = StudentAgentMessage(session_id=session.id, role="assistant", content="")
    db.add(assistant_message)
    db.commit()
    db.refresh(assistant_message)

    full_content = ""
    async for event_name, data in run_agent_loop(
        db, identity, session, user_message, assistant_message,
        model, messages, openai_tools, registry, attachments, reasoning_effort,
        max_iterations, permission_mode, config.temperature, config.max_tokens,
    ):
        if event_name == "message.delta":
            full_content += str(data.get("delta", ""))
        yield dumps_event(event_name, data)

    if not full_content.strip():
        full_content = _configured_fallback_answer(config, content)
        yield dumps_event("message.delta", {"message_id": assistant_message.id, "delta": full_content})

    assistant_message.content = full_content
    session.updated_at = utcnow()
    db.commit()
    yield dumps_event("message.completed", {"message_id": assistant_message.id})
    yield dumps_event("done", {"session_id": session.id})


def _claim_message_attachments(
    db: Session,
    identity: AuthIdentity,
    session: StudentAgentSession,
    message: StudentAgentMessage,
    attachment_ids: list[int],
) -> list[StudentAgentAttachment]:
    if not attachment_ids:
        return []
    rows = list(
        db.scalars(
            select(StudentAgentAttachment).where(
                StudentAgentAttachment.id.in_(attachment_ids),
                StudentAgentAttachment.tenant_id == identity.tenant_id,
                StudentAgentAttachment.student_id == identity.user_id,
                StudentAgentAttachment.session_id == session.id,
            )
        ).all()
    )
    found = {row.id for row in rows}
    missing = [item for item in attachment_ids if item not in found]
    if missing:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="附件不存在或不属于当前会话")
    for row in rows:
        row.message_id = message.id
    db.commit()
    return rows



def _tool_start_label(tool: ToolDefinition, arguments: dict[str, Any]) -> str:
    """Human-readable 'in progress' label shown in the activity chip."""
    kind = tool.metadata.get("kind") or tool.source
    if kind == "profile":
        return "正在读取学生档案…"
    if kind == "resume":
        return "正在读取简历材料…"
    if kind == "file":
        return "正在解析上传附件…"
    if kind == "job":
        kw = str(arguments.get("keyword") or "")[:20]
        return f"正在检索岗位库{('：' + kw) if kw else '…'}"
    if kind == "knowledge":
        q = str(arguments.get("query") or "")[:20]
        return f"正在检索知识库{('：' + q) if q else '…'}"
    if kind == "context":
        return "正在回溯会话上下文…"
    if kind == "subagent":
        key = str(arguments.get("agent_key") or tool.name)
        return f"正在调用子智能体：{key}…"
    if kind == "notification":
        return "正在准备通知…"
    if tool.source == "skill":
        return f"正在调用 Skill：{tool.metadata.get('name') or tool.name}…"
    if tool.source == "mcp":
        return "正在探索 MCP 工具…"
    return f"正在执行 {tool.name}…"



def _invoke_skill(tool: ToolDefinition, arguments: dict[str, Any]) -> dict[str, Any]:
    skill_name = str(tool.metadata.get("name") or tool.name)
    return {
        "status": "completed",
        "tool": tool.name,
        "skill_slug": tool.metadata.get("slug"),
        "summary": f"已调用 Skill：{skill_name}，处理「{str(arguments.get('task') or '')[:30]}」。",
        # Skill 是「渐进式披露」的操作手册，调用时应把完整正文加载进上下文（不是 1600 字的缩略）
        "skill_content": str(tool.metadata.get("content") or "")[:12000],
        "description": tool.description,
    }


def _analyze_uploaded_files(attachments: list[StudentAgentAttachment]) -> dict[str, Any]:
    if not attachments:
        return {
            "status": "failed",
            "tool": "analyze_uploaded_file",
            "summary": "本轮消息没有可分析的附件。",
        }
    file_summaries = []
    for attachment in attachments:
        text = (attachment.extracted_text or "").strip()
        excerpt = text[:700] if text else "未提取到文本内容"
        file_summaries.append(
            {
                "id": attachment.id,
                "name": attachment.original_name,
                "content_type": attachment.content_type,
                "file_ext": attachment.file_ext,
                "file_size": attachment.file_size,
                "excerpt": excerpt,
            }
        )
    names = "、".join(item["name"] for item in file_summaries)
    has_image = any(attachment.content_type.startswith("image/") for attachment in attachments)
    image_note = (
        " 图片已解析，如模型支持视觉输入将直传。"
        if has_image
        else ""
    )
    return {
        "status": "completed",
        "tool": "analyze_uploaded_file",
        "summary": f"已解析附件：{names}。{image_note}",
        "attachments": file_summaries,
    }



# ── Web tools (Jina Reader) ───────────────────────────────────────────────────


def _read_webpage_tool(args: dict[str, Any]) -> dict[str, Any]:
    """通过 Jina Reader 读取网页内容，返回 Markdown。"""
    url = str(args.get("url") or "").strip()
    if not url:
        return {"status": "failed", "tool": "read_webpage", "summary": "缺少 url 参数。"}
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    max_length = int(args.get("max_length") or 5000)

    try:
        jina_url = f"https://r.jina.ai/{url}"
        with httpx.Client(timeout=15, follow_redirects=True) as client:
            resp = client.get(jina_url, headers={"Accept": "text/plain"})
            resp.raise_for_status()
            content = resp.text[:max_length]
        return {
            "status": "completed",
            "tool": "read_webpage",
            "summary": f"已读取网页内容（{len(content)} 字符）。",
            "url": url,
            "content": content,
        }
    except httpx.TimeoutException:
        return {"status": "failed", "tool": "read_webpage", "summary": f"读取超时：{url}"}
    except httpx.HTTPStatusError as exc:
        return {"status": "failed", "tool": "read_webpage", "summary": f"HTTP {exc.response.status_code}：{url}"}
    except Exception as exc:
        logger.warning("read_webpage 失败: %s", exc)
        return {"status": "failed", "tool": "read_webpage", "summary": f"读取失败：{exc}"}


def _web_search_tool(args: dict[str, Any]) -> dict[str, Any]:
    """联网搜索关键词。优先用 Jina Search API（需 JINA_API_KEY），否则通过 Jina Reader 抓 DuckDuckGo 结果页。"""
    import os
    from urllib.parse import quote_plus

    query = str(args.get("query") or "").strip()
    if not query:
        return {"status": "failed", "tool": "web_search", "summary": "缺少 query 参数。"}

    jina_key = os.environ.get("JINA_API_KEY", "")

    # 方式一：Jina Search API（需 API Key）
    if jina_key:
        try:
            with httpx.Client(timeout=15, follow_redirects=True) as client:
                resp = client.get(
                    f"https://s.jina.ai/{query}",
                    headers={"Accept": "text/plain", "Authorization": f"Bearer {jina_key}"},
                )
                resp.raise_for_status()
                content = resp.text[:8000]
            return {
                "status": "completed",
                "tool": "web_search",
                "summary": f"已搜索「{query}」。",
                "query": query,
                "content": content,
            }
        except Exception:
            logger.debug("Jina Search API 调用失败，回退到 DuckDuckGo")
            pass  # 回退到方式二

    # 方式二：通过 Jina Reader 抓 DuckDuckGo 搜索结果页（免费）
    try:
        ddg_url = f"https://lite.duckduckgo.com/lite/?q={quote_plus(query)}"
        jina_url = f"https://r.jina.ai/{ddg_url}"
        with httpx.Client(timeout=15, follow_redirects=True) as client:
            resp = client.get(jina_url, headers={"Accept": "text/plain"})
            resp.raise_for_status()
            content = resp.text[:8000]
        if not content.strip():
            raise ValueError("空内容")
        return {
            "status": "completed",
            "tool": "web_search",
            "summary": f"已搜索「{query}」（DuckDuckGo）。",
            "query": query,
            "content": content,
        }
    except Exception:
        logger.warning("web_search 所有方式均失败: query=%s", query)
        pass  # 回退到方式三

    # 方式三：回退到 read_webpage，让模型用已知 URL 自行补充
    return {
        "status": "partial",
        "tool": "web_search",
        "summary": (
            f"无法直接搜索「{query}」。建议：请学生提供具体网址，使用 read_webpage 工具读取；"
            "或在回复中引导学生自行搜索后粘贴链接。"
        ),
        "query": query,
        "fallback_hint": "read_webpage",
    }



def _query_student_profile(db: Session, identity: AuthIdentity) -> dict[str, Any]:
    student = db.get(StudentUser, identity.user_id)
    if not student:
        return {"status": "failed", "tool": "query_student_profile", "summary": "没有找到学生档案。"}
    profile = {
        "name": student.name or "同学",
        "email": student.email,
        "college": student.college,
        "major": student.major,
        "grade": student.grade,
    }
    visible = "，".join(f"{k}={v}" for k, v in profile.items() if v)
    return {
        "status": "completed",
        "tool": "query_student_profile",
        "summary": f"已读取学生档案：{visible or '暂无详细背景'}。",
        "profile": profile,
    }


def _get_session_context(db: Session, session: StudentAgentSession, limit: int) -> dict[str, Any]:
    messages = list(
        db.scalars(
            select(StudentAgentMessage)
            .where(StudentAgentMessage.session_id == session.id)
            .order_by(StudentAgentMessage.id.desc())
            .limit(max(1, min(limit, 20)))
        ).all()
    )
    context = [{"role": item.role, "content": item.content[:500]} for item in reversed(messages)]
    return {
        "status": "completed",
        "tool": "get_session_context",
        "summary": f"已回溯 {len(context)} 条会话记录。",
        "messages": context,
    }




def _select_chat_model(
    db: Session,
    tenant_id: int,
    requested_model_id: Optional[int],
    allowed_capabilities: tuple[str, ...] = CHAT_CAPABLE_CAPABILITIES,
) -> Optional[ModelConfig]:
    if requested_model_id:
        model = db.get(ModelConfig, requested_model_id)
        if (
            model
            and model.tenant_id == tenant_id
            and not model.is_deleted
            and model.open_to_student
            and model.capability in allowed_capabilities
            and model.status == "active"
        ):
            return model
        return None

    config = get_or_create_master_config(db, tenant_id)
    if config.model_id:
        model = db.get(ModelConfig, config.model_id)
        if (
            model
            and model.tenant_id == tenant_id
            and not model.is_deleted
            and model.open_to_student
            and model.capability in allowed_capabilities
            and model.status == "active"
        ):
            return model
    return db.scalar(
        select(ModelConfig)
        .where(
            ModelConfig.tenant_id == tenant_id,
            ModelConfig.is_deleted.is_(False),
            ModelConfig.open_to_student.is_(True),
            ModelConfig.capability.in_(allowed_capabilities),
            ModelConfig.status == "active",
        )
        .order_by(ModelConfig.id.asc())
    )




def _effort_instruction(reasoning_effort: str) -> str:
    labels = {
        "low": "低。快速响应，给出简洁可执行建议。",
        "medium": "中。平衡速度和质量，覆盖关键依据与下一步。",
        "high": "高。充分分析，补齐风险和细节。",
        "xhigh": "超高。系统拆解、多角度验证，给出完整行动计划。",
    }
    return labels.get(reasoning_effort, labels["medium"])


def _attachment_prompt_text(
    attachments: list[StudentAgentAttachment],
    inline_images: bool = False,
) -> str:
    if not attachments:
        return "无附件。"
    chunks: list[str] = []
    for attachment in attachments:
        is_image = attachment.content_type.startswith("image/")
        if is_image and inline_images:
            # The raw image is attached inline to this same message — tell the
            # model to look at it directly instead of relying on metadata.
            body = "（图片已随本条消息一并传入，请直接观察图片内容进行分析，不要回答“只能看到元数据”。）"
        else:
            extracted = (attachment.extracted_text or "").strip()[:3000]
            body = f"提取内容:\n{extracted or '未提取到文本内容。'}"
        chunks.append(
            "\n".join(
                [
                    f"附件 {attachment.id}: {attachment.original_name}",
                    f"类型: {attachment.content_type}, 大小: {attachment.file_size} bytes",
                    body,
                ]
            )
        )
    return "\n\n".join(chunks)


def _attachment_image_parts(attachments: list[StudentAgentAttachment]) -> list[dict[str, Any]]:
    parts: list[dict[str, Any]] = []
    for attachment in attachments:
        if not attachment.content_type.startswith("image/"):
            continue
        if attachment.file_size > 8_000_000:
            continue
        path = Path(attachment.stored_path)
        if not path.exists():
            continue
        data_url = f"data:{attachment.content_type};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"
        parts.append({"type": "image_url", "image_url": {"url": data_url}})
    return parts[:4]


def _has_image_attachments(attachments: list[StudentAgentAttachment]) -> bool:
    return any(attachment.content_type.startswith("image/") for attachment in attachments)


def _supports_reasoning_effort(model: ModelConfig) -> bool:
    haystack = f"{model.provider} {model.model_identifier} {model.protocols}".lower()
    return any(token in haystack for token in ["openai", "o1", "o3", "o4", "gpt-5"])


def _supports_image_input(model: ModelConfig) -> bool:
    """Assume chat models are multimodal and pass images through by default.

    A keyword allowlist is too fragile — new/custom multimodal model names
    (e.g. mimo-v2.5) get silently dropped. Instead we attempt image passthrough
    for every chat model and only skip it for modalities that obviously cannot
    handle images (embedding / rerank / speech). If a genuinely text-only chat
    model rejects the request, the stream degrades to the text fallback answer.
    """
    haystack = f"{model.provider} {model.model_identifier} {model.protocols} {model.capability}".lower()
    non_visual_markers = [
        "embedding", "embed", "rerank", "reranker",
        "whisper", "tts", "speech", "audio", "voice",
        "moderation", "guard",
    ]
    return not any(token in haystack for token in non_visual_markers)


def _fallback_answer(user_text: str, observations: list[RuntimeObservation]) -> str:
    done = [item.summary for item in observations]
    prefix = "\n".join(f"- {item}" for item in done) if done else "- 已记录你的需求。"
    return (
        "我已整理好以下上下文：\n\n"
        f"{prefix}\n\n"
        "接下来你可以继续描述需求，例如上传简历、告诉我目标岗位，"
        "我会调用合适的 Skill 或子智能体进一步处理，并把结果汇总给你。"
    )


def _configured_fallback_answer(config: Any, user_text: str) -> str:
    mode = str(getattr(config, "fallback_mode", "") or "direct_answer").lower()
    custom_message = str(getattr(config, "fallback_message", "") or "").strip()
    if mode == "guide_message":
        return custom_message or (
            "这次我没能完成处理。你可以补充目标岗位、简历或具体问题后重试，"
            "我会重新选择合适的工具继续处理。"
        )
    if mode == "error":
        return custom_message or "主智能体暂时无法完成本次请求，请稍后重试。"
    return _fallback_answer(user_text, [])


# ══════════════════════════════════════════════════════════════════════════════
# Agentic Loop（Model + Harness）—— Model 只提议工具，Harness 负责执行/校验/审计
# ══════════════════════════════════════════════════════════════════════════════

# 仅暴露 Harness 能够「诚实兑现」的工具。会编造结果的占位工具（岗位库 / 知识库 /
# 子智能体 / MCP）在内核稳定前一律不进工具池——对应准则「禁止编造经营结果」。
ACTIVE_BUILTIN_TOOL_NAMES = (
    "query_student_profile",
    "read_resume",
    "analyze_uploaded_file",
    "get_session_context",
    "export_resume_pdf",
    "read_webpage",
    "web_search",
    "generate_resume_data",
    "optimize_resume_data",
    "update_resume_data",
)


def assemble_active_tools(db: Session, identity: AuthIdentity) -> list[ToolDefinition]:
    """组装工具池：内置工具白名单 + 已启用的 Skill + 已启用的子智能体（每条路由一个工具）。"""
    pool: dict[str, ToolDefinition] = {}
    by_name = {tool.name: tool for tool in BUILTIN_TOOLS}
    for name in ACTIVE_BUILTIN_TOOL_NAMES:
        tool = by_name.get(name)
        if tool:
            pool[name] = tool

    for skill in list_skills(db, include_disabled=False):
        data = serialize_skill(skill)
        name = "skill__" + _tool_safe_name(str(data["slug"]))
        if name in pool:
            continue
        pool[name] = ToolDefinition(
            name=name,
            description=str(data.get("description") or data.get("name") or "Skill 工具"),
            source="skill",
            priority=500,
            input_schema={
                "type": "object",
                "properties": {"task": {"type": "string", "description": "交给该 Skill 处理的具体任务。"}},
                "required": ["task"],
            },
            metadata=data,
        )

    # 设计决策（2026-06）：主智能体不再调用子智能体。任务型能力（简历优化/岗位匹配）做成
    # Skill 由主智能体编排；沉浸型人格（AI 面试官/职业规划师/岗位推荐师）放在「智能体广场」，
    # 由学生直接进入多轮对话——把有状态人格压成一次性工具调用会毁掉其多轮体验。
    return sorted(pool.values(), key=lambda item: (-item.priority, item.name))


def _build_openai_tools(tool_defs: list[ToolDefinition]) -> list[dict[str, Any]]:
    return [
        {
            "type": "function",
            "function": {
                "name": tool.name,
                "description": (tool.description or "")[:1024],
                "parameters": tool.input_schema or {"type": "object", "properties": {}},
            },
        }
        for tool in tool_defs
    ]


def _harness_system_prompt(config: Any, reasoning_effort: str) -> str:
    persona = (getattr(config, "system_prompt", None) or DEFAULT_SYSTEM_PROMPT).strip()
    effort = _effort_instruction(reasoning_effort)
    rules = (
        "\n\n## 运行机制（Harness 管控，必须遵守）\n"
        "- 你只负责理解需求、规划步骤、选择工具、综合结果；工具的执行、校验与权限由 Harness 负责，你无需也不能自行控制循环。\n"
        "- 工具纪律：只能调用系统提供给你的工具，并严格按其参数 schema 传参；不要臆测不存在的能力，也不要伪造任何工具的返回结果。\n"
        "- 反幻觉铁律：禁止编造学生的简历内容、经历、项目、岗位、公司或任何数据。没有依据时如实说明，并向学生索取材料。\n"
        "- 简历相关：在给出任何简历修改建议或生成简历之前，必须先调用 read_resume 读取学生的真实简历；"
        "若 read_resume 返回没有简历，请直接告知并引导学生到『个人中心—我的简历』上传，绝不虚构内容。\n"
        "- AI 简历制作流程：先调用 query_student_profile 获取学生信息，再请学生提供目标岗位 JD，"
        "然后调用 generate_resume_data 生成结构化简历；工具返回 editor_url 时，"
        "用 Markdown 链接 [点击查看并编辑简历](editor_url) 呈现给学生。\n"
        "- 简历优化流程：先调用 read_resume 读取学生已有简历，再请学生提供目标 JD，"
        "然后调用 optimize_resume_data 生成优化版本；工具返回 editor_url 时，"
        "用 Markdown 链接 [点击查看优化后的简历](editor_url) 呈现给学生。\n"
        "- 修改已有在线简历：调用 update_resume_data（需提供 resume_id），工具返回 editor_url 后用链接呈现。\n"
        "- 生成可下载简历：当学生需要『修改好的 / 可下载的简历』时，先基于真实简历完成改写，再调用 "
        "export_resume_pdf（传入完整的 Markdown 简历正文）生成 PDF，然后把工具返回的 download_url 以 "
        "Markdown 链接形式给学生，例如：[点击下载优化后的简历](下载链接)。\n"
        "- 沉浸式专家：当学生需要『模拟面试 / AI 面试官』『职业规划咨询』『岗位推荐』等多轮、有人格的沉浸体验时，"
        "你不要自己扮演，而是引导学生前往『智能体广场』进入对应的专属智能体（那里才是多轮对话的入口）。\n"
        "- 联网工具：当学生发来 URL 链接或需要查看网页内容时，调用 read_webpage 读取；"
        "当需要搜索公司信息、行业动态等实时数据时，调用 web_search 搜索。"
        "如果搜索失败，引导学生自行搜索后粘贴链接，再用 read_webpage 读取。\n"
        "- 输出规范：使用 Markdown，先结论后步骤；不要输出工具调用的原始 JSON、tool_call 或隐藏推理过程。\n"
        f"- 推理强度：{effort}"
    )
    return persona + rules


def _build_initial_messages(
    db: Session,
    identity: AuthIdentity,
    session: StudentAgentSession,
    user_text: str,
    reasoning_effort: str,
    model: ModelConfig,
    attachments: list[StudentAgentAttachment],
    config: Any,
) -> list[dict[str, Any]]:
    messages: list[dict[str, Any]] = [
        {"role": "system", "content": _harness_system_prompt(config, reasoning_effort)}
    ]

    # 取最近 24 条（desc 后反转为 asc），最后一条即本轮 user 消息，丢弃避免重复。
    history_rows = list(
        db.scalars(
            select(StudentAgentMessage)
            .where(StudentAgentMessage.session_id == session.id)
            .order_by(StudentAgentMessage.id.desc())
            .limit(24)
        ).all()
    )
    history_rows.reverse()
    for msg in history_rows[:-1]:
        if msg.role not in ("user", "assistant"):
            continue
        text = msg.content
        if len(text) > 4000:
            text = text[:4000] + "\n…[已截断]"
        messages.append({"role": msg.role, "content": text})

    inline_images = _has_image_attachments(attachments) and _supports_image_input(model)
    parts = [user_text]
    if attachments:
        parts.append("\n---\n**本轮附件**\n" + _attachment_prompt_text(attachments, inline_images))
    current_text = "\n".join(parts)

    if inline_images:
        messages.append(
            {"role": "user", "content": [{"type": "text", "text": current_text}, *_attachment_image_parts(attachments)]}
        )
    else:
        messages.append({"role": "user", "content": current_text})
    return messages


async def run_agent_loop(
    db: Session,
    identity: AuthIdentity,
    session: StudentAgentSession,
    user_message: StudentAgentMessage,
    assistant_message: StudentAgentMessage,
    model: ModelConfig,
    messages: list[dict[str, Any]],
    openai_tools: list[dict[str, Any]],
    registry: dict[str, ToolDefinition],
    attachments: list[StudentAgentAttachment],
    reasoning_effort: str,
    max_iterations: int,
    permission_mode: str = "ask",
    temperature: Optional[float] = None,
    max_tokens: Optional[int] = None,
) -> AsyncIterator[tuple[str, dict[str, Any]]]:
    """The harness-owned ReAct loop. Yields (sse_event_name, data) tuples."""
    assistant_id = assistant_message.id
    deadline = time.monotonic() + 300  # 5 分钟总超时
    logger.info("agent_loop 开始 session=%s model=%s max_iter=%s", session.id, model.model_identifier, max_iterations)

    for iteration in range(max_iterations):
        if time.monotonic() > deadline:
            logger.warning("agent_loop 超时 session=%s iteration=%s", session.id, iteration)
            yield "message.delta", {"message_id": assistant_id, "delta": "\n\n[回复超时，请重试]"}
            break
        turn_content = ""
        turn_tool_calls: list[dict[str, Any]] = []
        turn_error = False

        async for kind, value in _stream_llm_turn(
            model, messages, openai_tools, reasoning_effort, temperature, max_tokens
        ):
            if kind == "delta":
                yield "message.delta", {"message_id": assistant_id, "delta": value}
            elif kind == "error":
                turn_error = True
            elif kind == "final":
                turn_content = value.get("content") or ""
                turn_tool_calls = value.get("tool_calls") or []

        # 模型不支持 tools（请求报错）→ 降级：去掉 tools 再要一次纯文本回答。
        if turn_error and not turn_tool_calls:
            async for kind, value in _stream_llm_turn(
                model, messages, [], reasoning_effort, temperature, max_tokens
            ):
                if kind == "delta":
                    yield "message.delta", {"message_id": assistant_id, "delta": value}
            return

        if not turn_tool_calls:
            return  # 最终回答已流式输出完毕

        # 规范每个 tool_call 的 id，保证 assistant 消息与 tool 结果一一对应。
        for i, tc in enumerate(turn_tool_calls):
            if not tc.get("id"):
                tc["id"] = f"call_{iteration}_{i}"

        messages.append(
            {
                "role": "assistant",
                "content": turn_content or "",
                "tool_calls": [
                    {
                        "id": tc["id"],
                        "type": "function",
                        "function": {"name": tc.get("name") or "", "arguments": tc.get("arguments") or "{}"},
                    }
                    for tc in turn_tool_calls
                ],
            }
        )

        for tc in turn_tool_calls:
            name = tc.get("name") or ""
            call_id = tc["id"]
            td = registry.get(name)
            args, argument_errors = parse_tool_arguments(
                tc.get("arguments"),
                td.input_schema if td else None,
            )
            activity_kind = (td.metadata.get("kind") if td else None) or (td.source if td else None) or "context"
            start_label = (
                f"工具「{name}」参数校验失败"
                if argument_errors
                else (_tool_start_label(td, args) if td else f"正在处理未知工具 {name}…")
            )

            # 四态权限裁决（Harness 管控）：未通过则不执行，回结构化结果让模型转而向学生说明。
            decision, deny_reason = _permission_decision(permission_mode, name, td)

            started = _save_activity(
                db, session, user_message,
                kind=str(activity_kind), name=name or "unknown",
                status_value="started",
                summary=deny_reason if decision != "allow" else start_label,
                detail={
                    "iteration": iteration + 1,
                    "tool_call_id": call_id,
                    "arguments": args,
                    "argument_errors": argument_errors,
                    "permission_mode": permission_mode,
                    "decision": decision,
                },
            )
            yield "activity.started", serialize_activity(started).model_dump(mode="json")

            if argument_errors:
                result = {
                    "status": "failed",
                    "tool": name,
                    "summary": "；".join(argument_errors),
                    "error_code": "invalid_tool_arguments",
                }
            elif decision == "allow":
                result = await _dispatch_tool(
                    db, identity, session, assistant_message, user_message.content, attachments, name, args, td
                )
            else:
                result = {"status": "failed", "tool": name, "summary": deny_reason, "permission": decision}
            result_detail = {
                **result,
                "iteration": iteration + 1,
                "tool_call_id": call_id,
                "arguments": args,
            }
            completed = _complete_activity(
                db, started,
                status_value=result.get("status", "completed"),
                summary=result.get("summary", ""),
                detail=result_detail,
            )
            event_name = "activity.completed" if result.get("status") == "completed" else "activity.failed"
            yield event_name, serialize_activity(completed).model_dump(mode="json")

            # 生成了可下载文件时，额外推一个事件供前端做下载入口（前端忽略也无害）。
            if result.get("download_url"):
                yield "attachment.created", {
                    "message_id": assistant_id,
                    "download_url": result["download_url"],
                    "filename": result.get("filename"),
                    "attachment_id": result.get("attachment_id"),
                }

            messages.append({"role": "tool", "tool_call_id": call_id, "content": _tool_result_for_model(result)})

    # 触顶 max_iterations —— 强制一次无工具的收尾回答，避免无限循环。
    async for kind, value in _stream_llm_turn(
        model, messages, [], reasoning_effort, temperature, max_tokens
    ):
        if kind == "delta":
            yield "message.delta", {"message_id": assistant_id, "delta": value}


async def _stream_llm_turn(
    model: ModelConfig,
    messages: list[dict[str, Any]],
    tools: list[dict[str, Any]],
    reasoning_effort: str,
    temperature: Optional[float] = None,
    max_tokens: Optional[int] = None,
) -> AsyncIterator[tuple[str, Any]]:
    """Single streaming turn. Yields ("delta", text) / ("error", msg) / ("final", dict)."""
    try:
        api_key = decrypt_api_key(model.api_key_cipher or "")
        payload: dict[str, Any] = {
            "model": model.model_identifier,
            "messages": messages,
            "temperature": temperature if temperature is not None else (
                model.default_temp if model.default_temp is not None else 0.7
            ),
            "max_tokens": max_tokens if max_tokens is not None else (model.max_output or 4096),
            "stream": True,
        }
        if tools:
            payload["tools"] = tools
            payload["tool_choice"] = "auto"
        if _supports_reasoning_effort(model):
            payload["reasoning_effort"] = "high" if reasoning_effort == "xhigh" else reasoning_effort

        tool_calls_acc: dict[int, dict[str, str]] = {}
        content_acc = ""
        finish: Optional[str] = None

        async with httpx.AsyncClient(
            timeout=httpx.Timeout(connect=10, read=model.timeout_sec or 60, write=30, pool=5)
        ) as client:
            async with client.stream(
                "POST",
                f"{model.base_url.rstrip('/')}/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json=payload,
            ) as response:
                response.raise_for_status()
                async for line in response.aiter_lines():
                    if not line.startswith("data:"):
                        continue
                    raw = line[5:].strip()
                    if raw == "[DONE]":
                        break
                    try:
                        obj = json.loads(raw)
                    except json.JSONDecodeError:
                        continue
                    choice = (obj.get("choices") or [{}])[0]
                    delta = choice.get("delta") or {}
                    piece = delta.get("content")
                    if piece:
                        content_acc += piece
                        yield "delta", piece
                    for tc in delta.get("tool_calls") or []:
                        idx = tc.get("index", 0) or 0
                        slot = tool_calls_acc.setdefault(idx, {"id": "", "name": "", "arguments": ""})
                        if tc.get("id"):
                            slot["id"] = tc["id"]
                        fn = tc.get("function") or {}
                        if fn.get("name"):
                            slot["name"] = fn["name"]
                        if fn.get("arguments"):
                            slot["arguments"] += fn["arguments"]
                    if choice.get("finish_reason"):
                        finish = choice["finish_reason"]
    except Exception as exc:  # noqa: BLE001 — surfaced to caller for graceful fallback
        logger.exception("LLM 流式调用失败")
        yield "error", str(exc)[:200]
        return

    ordered = [tool_calls_acc[key] for key in sorted(tool_calls_acc.keys())]
    yield "final", {"content": content_acc, "tool_calls": ordered, "finish_reason": finish}


# ── Harness tool dispatch ───────────────────────────────────────────────────────


async def _dispatch_tool(
    db: Session,
    identity: AuthIdentity,
    session: StudentAgentSession,
    assistant_message: StudentAgentMessage,
    user_text: str,
    attachments: list[StudentAgentAttachment],
    name: str,
    args: dict[str, Any],
    td: Optional[ToolDefinition],
) -> dict[str, Any]:
    # 未知工具：返回结构化错误让模型自我纠正，而不是崩溃。
    if td is None:
        return {"status": "failed", "tool": name, "summary": f"未知工具「{name}」，已忽略。请只调用系统提供的工具。"}
    if td.source == "skill":
        return _invoke_skill(td, args)
    if name == "query_student_profile":
        return _query_student_profile(db, identity)
    if name == "read_resume":
        return _read_resume_tool(db, identity, session, attachments)
    if name == "analyze_uploaded_file":
        return _analyze_uploaded_files(attachments)
    if name == "get_session_context":
        return _get_session_context(db, session, int(args.get("limit") or 8))
    if name == "export_resume_pdf":
        return _export_resume_pdf_tool(db, identity, session, assistant_message, args)
    if name == "read_webpage":
        return _read_webpage_tool(args)
    if name == "web_search":
        return _web_search_tool(args)
    if name == "generate_resume_data":
        return _generate_resume_data_tool(db, identity, args)
    if name == "optimize_resume_data":
        return _optimize_resume_data_tool(db, identity, args)
    if name == "update_resume_data":
        return _update_resume_data_tool(db, identity, args)
    return {"status": "failed", "tool": name, "summary": f"工具 {name} 暂未接入执行器。"}


_TOOL_RESULT_KEYS_TO_STRIP = {"tool", "status", "iteration", "tool_call_id", "arguments"}


def _tool_result_for_model(result: dict[str, Any]) -> str:
    """序列化工具结果发给模型，去掉内部元数据字段以节省 context window。"""
    filtered = {k: v for k, v in result.items() if k not in _TOOL_RESULT_KEYS_TO_STRIP}
    try:
        text = json.dumps(filtered, ensure_ascii=False)
    except (TypeError, ValueError):
        text = str(filtered)
    return text[:6000]


# ── 四态权限裁决（allow / ask / deny）────────────────────────────────────────────


def _permission_decision(mode: str, name: str, td: Optional[ToolDefinition]) -> tuple[str, str]:
    """根据主智能体配置的 permission_mode 与工具风险等级裁决是否执行。

    返回 (decision, reason)，decision ∈ {"allow", "ask", "deny"}。
    - auto：除被 deny 标记外，一律放行；
    - ask（默认）：低风险放行，需确认的高风险动作暂缓（当前工具池无此类，预留给投递/发信等）；
    - strict：仅放行平台内置安全工具，Skill 与子智能体一律拒绝。
    未知工具交给 _dispatch_tool 返回结构化错误，这里直接放行。
    """
    if td is None:
        return "allow", ""

    risk_raw = str(td.metadata.get("risk", "")).lower()
    if risk_raw in ("deny", "high", "critical"):
        risk = "deny"
    elif risk_raw in ("confirm", "ask", "medium"):
        risk = "confirm"
    else:
        risk = "allow"

    strict_ok = td.source == "builtin"  # 仅平台内置工具进入 strict 白名单

    if risk == "deny":
        return "deny", f"工具「{name}」已被 Harness 禁用，拒绝执行。"
    if mode == "strict" and not strict_ok:
        return (
            "deny",
            f"当前为 strict 权限模式，仅允许平台内置安全工具，已拒绝调用「{name}」。"
            "请改用内置工具，或如实告知学生该能力当前不可用。",
        )
    if mode == "ask" and risk == "confirm":
        return (
            "ask",
            f"工具「{name}」属于需要学生确认的操作，当前未获确认，已暂缓。"
            "请先向学生说明将要执行的动作并征得同意。",
        )
    return "allow", ""


# ── Resume tools ────────────────────────────────────────────────────────────────


def _ensure_attachment_text(db: Session, attachment: StudentAgentAttachment) -> str:
    """Lazily extract & persist text for an attachment that has none."""
    existing = (attachment.extracted_text or "").strip()
    if existing:
        return existing
    path = Path(attachment.stored_path)
    if not path.exists():
        return ""
    try:
        if attachment.file_ext == "pdf":
            text = _extract_pdf_text(path)
        else:
            text = _extract_attachment_text(path, attachment.content_type, attachment.file_ext)
    except Exception:
        return ""
    if text and text.strip():
        attachment.extracted_text = text
        try:
            db.commit()
        except Exception:
            db.rollback()
        return text
    return ""


def _rich_text_to_lines(html: str) -> list[str]:
    """Convert HTML rich text to plain text lines (mirrors frontend richTextToLines)."""
    if not html:
        return []
    text = _re.sub(r"<br\s*/?>", "\n", html, flags=_re.IGNORECASE)
    text = _re.sub(r"<li[^>]*>", "\n", text, flags=_re.IGNORECASE)
    text = _re.sub(r"</(p|div|section|li|ul|ol|h[1-6])>", "\n", text, flags=_re.IGNORECASE)
    text = _re.sub(r"<[^>]+>", "", text)
    for entity, char in [("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"), ("&quot;", '"'), ("&#39;", "'")]:
        text = text.replace(entity, char)
    return [ln.strip() for ln in text.split("\n") if ln.strip()]


def _ta_to_list(text: Any) -> str:
    """Convert newline-separated plain text to <ul><li>…</li></ul> HTML."""
    lines = [ln.strip() for ln in str(text or "").splitlines() if ln.strip()]
    if not lines:
        return ""
    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return "<ul>" + "".join(f"<li>{esc(ln)}</li>" for ln in lines) + "</ul>"


def _ta_to_para(text: Any) -> str:
    """Convert newline-separated plain text to <p>…</p> HTML blocks."""
    lines = [ln.strip() for ln in str(text or "").splitlines() if ln.strip()]
    if not lines:
        return ""
    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return "".join(f"<p>{esc(ln)}</p>" for ln in lines)


def _structured_resume_to_text(row: StudentResume) -> str:
    """Convert a StudentResume row to readable plain text for the AI model."""
    try:
        data = json.loads(row.data_json or "{}")
    except Exception:
        data = {}
    basic = data.get("basic") or {}
    parts: list[str] = []
    for label, key in [("姓名", "name"), ("目标职位", "title"), ("邮箱", "email"), ("电话", "phone"), ("地址", "location"), ("生日", "birthDate")]:
        val = str(basic.get(key) or "").strip()
        if val:
            parts.append(f"{label}: {val}")
    skill_lines = _rich_text_to_lines(data.get("skillContent") or "")
    if skill_lines:
        parts.append("\n专业技能:")
        parts.extend(f"- {ln}" for ln in skill_lines)
    for exp in (data.get("experience") or []):
        if exp.get("visible") is False:
            continue
        header = " | ".join(v for v in [exp.get("company"), exp.get("position"), exp.get("date")] if v)
        if header:
            parts.append(f"\n工作经历: {header}")
        parts.extend(f"- {ln}" for ln in _rich_text_to_lines(exp.get("details") or ""))
    for proj in (data.get("projects") or []):
        if proj.get("visible") is False:
            continue
        header = " | ".join(v for v in [proj.get("name"), proj.get("role"), proj.get("date")] if v)
        if header:
            parts.append(f"\n项目经历: {header}")
        parts.extend(f"- {ln}" for ln in _rich_text_to_lines(proj.get("description") or ""))
    for edu in (data.get("education") or []):
        if edu.get("visible") is False:
            continue
        header = " | ".join(v for v in [edu.get("school"), edu.get("major"), edu.get("degree"), f"{edu.get('startDate', '')}-{edu.get('endDate', '')}"] if v)
        if header:
            parts.append(f"\n教育经历: {header}")
        parts.extend(f"- {ln}" for ln in _rich_text_to_lines(edu.get("description") or ""))
    eval_lines = _rich_text_to_lines(data.get("selfEvaluationContent") or "")
    if eval_lines:
        parts.append("\n自我评价:")
        parts.extend(eval_lines)
    return "\n".join(parts)


def _read_resume_tool(
    db: Session,
    identity: AuthIdentity,
    session: StudentAgentSession,
    attachments: list[StudentAgentAttachment],
) -> dict[str, Any]:
    """Read the student's resume — structured online resumes (visibility=True) first,
    then this turn's uploads, then profile-level PDF attachments."""
    resumes: list[dict[str, Any]] = []
    seen_att: set[int] = set()

    # 1. 在线结构化简历（visibility=True）
    structured_rows = list(
        db.scalars(
            select(StudentResume)
            .where(
                StudentResume.tenant_id == identity.tenant_id,
                StudentResume.student_id == identity.user_id,
                StudentResume.visibility.is_(True),
            )
            .order_by(StudentResume.updated_at.desc())
            .limit(3)
        ).all()
    )
    for row in structured_rows:
        text = _structured_resume_to_text(row)
        if text.strip():
            resumes.append({"source": "在线简历", "name": row.title, "resume_id": row.id, "excerpt": text[:4000]})

    # 2. 本轮上传附件
    for att in attachments:
        seen_att.add(att.id)
        text = _ensure_attachment_text(db, att)
        if text:
            resumes.append({"source": "本轮上传", "name": att.original_name, "excerpt": text[:3000]})

    # 3. 个人中心 PDF 附件
    pdf_rows = list(
        db.scalars(
            select(StudentAgentAttachment)
            .where(
                StudentAgentAttachment.tenant_id == identity.tenant_id,
                StudentAgentAttachment.student_id == identity.user_id,
                StudentAgentAttachment.file_ext == "pdf",
            )
            .order_by(StudentAgentAttachment.created_at.desc())
            .limit(3)
        ).all()
    )
    for row in pdf_rows:
        if row.id in seen_att:
            continue
        text = _ensure_attachment_text(db, row)
        if text:
            resumes.append({"source": "个人中心", "name": row.original_name, "excerpt": text[:3000]})

    if not resumes:
        return {
            "status": "completed",
            "tool": "read_resume",
            "summary": "未找到简历：学生还没有上传简历，也没有开启『智能体可读取』的在线简历。",
            "resumes": [],
        }
    names = "、".join(item["name"] for item in resumes[:4])
    return {"status": "completed", "tool": "read_resume", "summary": f"已读取简历：{names}", "resumes": resumes[:4]}


_MAX_RESUMES = 5
_VALID_TEMPLATE_IDS = {"classic", "modern", "elegant"}
_DEFAULT_GLOBAL_SETTINGS = {
    "classic": {
        "themeColor": "#000000",
        "fontFamily": '"Alibaba PuHuiTi", sans-serif',
        "baseFontSize": 16,
        "pagePadding": 32,
        "lineHeight": 1.5,
        "sectionSpacing": 16,
        "paragraphSpacing": 12,
        "headerSize": 18,
        "subheaderSize": 16,
        "useIconMode": True,
        "centerSubtitle": True,
    },
    "modern": {
        "themeColor": "#000000",
        "fontFamily": '"Alibaba PuHuiTi", sans-serif',
        "baseFontSize": 16,
        "pagePadding": 0,
        "lineHeight": 1.5,
        "sectionSpacing": 8,
        "paragraphSpacing": 4,
        "headerSize": 18,
        "subheaderSize": 16,
        "useIconMode": True,
        "centerSubtitle": True,
    },
    "elegant": {
        "themeColor": "#18181b",
        "fontFamily": '"Alibaba PuHuiTi", sans-serif',
        "baseFontSize": 16,
        "pagePadding": 32,
        "lineHeight": 1.5,
        "sectionSpacing": 28,
        "paragraphSpacing": 18,
        "headerSize": 20,
        "subheaderSize": 16,
        "useIconMode": True,
        "centerSubtitle": True,
    },
}
_DEFAULT_MENU_SECTIONS = [
    {"id": "basic", "title": "基本信息", "icon": "👤", "enabled": True, "order": 0},
    {"id": "skills", "title": "专业技能", "icon": "⚡", "enabled": True, "order": 1},
    {"id": "experience", "title": "工作经历", "icon": "💼", "enabled": True, "order": 2},
    {"id": "projects", "title": "项目经历", "icon": "🚀", "enabled": True, "order": 3},
    {"id": "education", "title": "教育经历", "icon": "🎓", "enabled": True, "order": 4},
    {"id": "selfEvaluation", "title": "自我评价", "icon": "📝", "enabled": True, "order": 5},
]
_DEFAULT_FIELD_ORDER = [
    {"id": "name", "key": "name", "label": "姓名", "type": "text", "visible": True},
    {"id": "title", "key": "title", "label": "职位", "type": "text", "visible": True},
    {"id": "birthDate", "key": "birthDate", "label": "生日", "type": "date", "visible": True},
    {"id": "employementStatus", "key": "employementStatus", "label": "状态", "type": "text", "visible": False},
    {"id": "email", "key": "email", "label": "邮箱", "type": "text", "visible": True},
    {"id": "phone", "key": "phone", "label": "电话", "type": "text", "visible": True},
    {"id": "location", "key": "location", "label": "地址", "type": "text", "visible": True},
]


def _build_resume_doc(args: dict[str, Any], student: Optional[Any], title: str, template_id: str) -> dict[str, Any]:
    """Build a full ResumeData-compatible document from AI-provided args."""
    basic_in = args.get("basic") or {}
    edu_in = args.get("education") or []
    exp_in = args.get("experience") or []
    proj_in = args.get("projects") or []

    def _edu(item: dict) -> dict:
        raw_desc = item.get("description") or ""
        return {
            "id": f"edu-{uuid.uuid4().hex[:8]}",
            "school": item.get("school") or "",
            "major": item.get("major") or "",
            "degree": item.get("degree") or "",
            "startDate": item.get("start_date") or item.get("startDate") or "",
            "endDate": item.get("end_date") or item.get("endDate") or "",
            "gpa": item.get("gpa") or "",
            "description": _ta_to_list(raw_desc) if raw_desc else "",
            "visible": True,
        }

    def _exp(item: dict) -> dict:
        raw = item.get("details") or item.get("description") or ""
        return {
            "id": f"exp-{uuid.uuid4().hex[:8]}",
            "company": item.get("company") or "",
            "position": item.get("position") or "",
            "date": item.get("date") or "",
            "details": _ta_to_list(raw) if raw else "",
            "visible": True,
        }

    def _proj(item: dict) -> dict:
        raw = item.get("description") or ""
        return {
            "id": f"proj-{uuid.uuid4().hex[:8]}",
            "name": item.get("name") or "",
            "role": item.get("role") or "",
            "date": item.get("date") or "",
            "description": _ta_to_list(raw) if raw else "",
            "visible": True,
            "link": "",
            "linkLabel": "",
        }

    basic = {
        "name": basic_in.get("name") or (getattr(student, "name", None) if student else None) or "",
        "title": basic_in.get("target_position") or basic_in.get("title") or "",
        "email": basic_in.get("email") or (getattr(student, "email", None) if student else None) or "",
        "phone": basic_in.get("phone") or (getattr(student, "phone", None) if student else None) or "",
        "location": basic_in.get("location") or (getattr(student, "college", None) if student else None) or "",
        "birthDate": basic_in.get("birth_date") or basic_in.get("birthDate") or "",
        "employementStatus": "",
        "photo": (getattr(student, "avatar_url", None) if student else None) or "",
        "icons": {"birthDate": "calendar", "employementStatus": "briefcase", "email": "mail", "phone": "phone", "location": "location"},
        "photoConfig": {"width": 90, "height": 120, "aspectRatio": "1:1", "borderRadius": "none", "customBorderRadius": 0, "visible": True},
        "fieldOrder": [dict(f) for f in _DEFAULT_FIELD_ORDER],
        "customFields": [],
        "githubKey": "",
        "githubUseName": "",
        "githubContributionsVisible": False,
    }

    skills_raw = args.get("skills") or ""
    self_eval_raw = args.get("self_evaluation") or ""

    return {
        "title": title,
        "templateId": template_id,
        "visibility": False,
        "basic": basic,
        "education": [_edu(item) for item in edu_in if isinstance(item, dict)],
        "experience": [_exp(item) for item in exp_in if isinstance(item, dict)],
        "projects": [_proj(item) for item in proj_in if isinstance(item, dict)],
        "certificates": [],
        "customData": {},
        "skillContent": _ta_to_list(skills_raw) if skills_raw else "",
        "selfEvaluationContent": _ta_to_para(self_eval_raw) if self_eval_raw else "",
        "activeSection": "basic",
        "draggingProjectId": None,
        "globalSettings": dict(_DEFAULT_GLOBAL_SETTINGS.get(template_id, _DEFAULT_GLOBAL_SETTINGS["classic"])),
        "menuSections": [dict(s) for s in _DEFAULT_MENU_SECTIONS],
    }


def _resume_count(db: Session, identity: AuthIdentity) -> int:
    return db.scalar(
        select(func.count(StudentResume.id)).where(
            StudentResume.student_id == identity.user_id,
            StudentResume.tenant_id == identity.tenant_id,
        )
    ) or 0


def _generate_resume_data_tool(db: Session, identity: AuthIdentity, args: dict[str, Any]) -> dict[str, Any]:
    if _resume_count(db, identity) >= _MAX_RESUMES:
        return {
            "status": "failed",
            "tool": "generate_resume_data",
            "summary": f"简历数量已达上限（{_MAX_RESUMES} 份），请先在『我的简历』中删除一份再生成。",
        }
    student = db.get(StudentUser, identity.user_id)
    title = str(args.get("title") or "AI 生成简历").strip()[:128] or "AI 生成简历"
    template_id = str(args.get("template_id") or "classic").strip()
    if template_id not in _VALID_TEMPLATE_IDS:
        template_id = "classic"
    doc = _build_resume_doc(args, student, title, template_id)
    row = StudentResume(
        tenant_id=identity.tenant_id,
        student_id=identity.user_id,
        title=title,
        template_id=template_id,
        visibility=False,
        data_json=json.dumps(doc, ensure_ascii=False),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return {
        "status": "completed",
        "tool": "generate_resume_data",
        "summary": f"简历《{title}》已生成，请点击链接进入编辑器查看并调整。",
        "resume_id": row.id,
        "editor_url": f"/student/resumes/{row.id}",
        "open_resume_editor": True,
    }


def _optimize_resume_data_tool(db: Session, identity: AuthIdentity, args: dict[str, Any]) -> dict[str, Any]:
    if _resume_count(db, identity) >= _MAX_RESUMES:
        return {
            "status": "failed",
            "tool": "optimize_resume_data",
            "summary": f"简历数量已达上限（{_MAX_RESUMES} 份），请先在『我的简历』中删除一份再优化。",
        }
    student = db.get(StudentUser, identity.user_id)
    title = str(args.get("title") or "优化版简历").strip()[:128] or "优化版简历"
    template_id = str(args.get("template_id") or "classic").strip()
    if template_id not in _VALID_TEMPLATE_IDS:
        # 如果来源简历有模板，则继承
        src_id = args.get("source_resume_id")
        if src_id:
            src_row = db.scalar(
                select(StudentResume).where(
                    StudentResume.id == int(src_id),
                    StudentResume.student_id == identity.user_id,
                    StudentResume.tenant_id == identity.tenant_id,
                )
            )
            template_id = (src_row.template_id if src_row else None) or "classic"
        else:
            template_id = "classic"
    doc = _build_resume_doc(args, student, title, template_id)
    row = StudentResume(
        tenant_id=identity.tenant_id,
        student_id=identity.user_id,
        title=title,
        template_id=template_id,
        visibility=False,
        data_json=json.dumps(doc, ensure_ascii=False),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return {
        "status": "completed",
        "tool": "optimize_resume_data",
        "summary": f"优化版简历《{title}》已生成，请点击链接进入编辑器查看并调整。",
        "resume_id": row.id,
        "editor_url": f"/student/resumes/{row.id}",
        "open_resume_editor": True,
    }


def _update_resume_data_tool(db: Session, identity: AuthIdentity, args: dict[str, Any]) -> dict[str, Any]:
    resume_id = args.get("resume_id")
    if not resume_id:
        return {"status": "failed", "tool": "update_resume_data", "summary": "缺少 resume_id 参数。"}
    row = db.scalar(
        select(StudentResume).where(
            StudentResume.id == int(resume_id),
            StudentResume.student_id == identity.user_id,
            StudentResume.tenant_id == identity.tenant_id,
        )
    )
    if not row:
        return {"status": "failed", "tool": "update_resume_data", "summary": f"简历 ID {resume_id} 不存在或无权限。"}

    try:
        existing = json.loads(row.data_json or "{}")
    except Exception:
        existing = {}

    # 合并标题和模板
    if args.get("title"):
        row.title = str(args["title"]).strip()[:128]
        existing["title"] = row.title
    if args.get("template_id") and args["template_id"] in _VALID_TEMPLATE_IDS:
        row.template_id = str(args["template_id"])
        existing["templateId"] = row.template_id

    student = db.get(StudentUser, identity.user_id)

    # 合并各字段（如果 AI 提供了就覆盖，否则保留原有）
    def _to_list_if_str(val: Any) -> Any:
        return _ta_to_list(val) if isinstance(val, str) else val

    if args.get("basic"):
        basic_in = args["basic"]
        existing_basic = existing.get("basic") or {}
        for key, ai_key in [("name", "name"), ("title", "target_position"), ("email", "email"), ("phone", "phone"), ("location", "location"), ("birthDate", "birth_date")]:
            val = basic_in.get(ai_key) or basic_in.get(key)
            if val:
                existing_basic[key] = val
        existing_basic.setdefault("title", basic_in.get("target_position") or basic_in.get("title") or existing_basic.get("title") or "")
        existing["basic"] = existing_basic

    if args.get("education") is not None:
        existing["education"] = [
            {
                "id": item.get("id") or f"edu-{uuid.uuid4().hex[:8]}",
                "school": item.get("school") or "",
                "major": item.get("major") or "",
                "degree": item.get("degree") or "",
                "startDate": item.get("start_date") or item.get("startDate") or "",
                "endDate": item.get("end_date") or item.get("endDate") or "",
                "gpa": item.get("gpa") or "",
                "description": _ta_to_list(item["description"]) if isinstance(item.get("description"), str) else item.get("description") or "",
                "visible": item.get("visible", True),
            }
            for item in args["education"] if isinstance(item, dict)
        ]

    if args.get("experience") is not None:
        existing["experience"] = [
            {
                "id": item.get("id") or f"exp-{uuid.uuid4().hex[:8]}",
                "company": item.get("company") or "",
                "position": item.get("position") or "",
                "date": item.get("date") or "",
                "details": _ta_to_list(item["details"]) if isinstance(item.get("details"), str) else item.get("details") or "",
                "visible": item.get("visible", True),
            }
            for item in args["experience"] if isinstance(item, dict)
        ]

    if args.get("projects") is not None:
        existing["projects"] = [
            {
                "id": item.get("id") or f"proj-{uuid.uuid4().hex[:8]}",
                "name": item.get("name") or "",
                "role": item.get("role") or "",
                "date": item.get("date") or "",
                "description": _ta_to_list(item["description"]) if isinstance(item.get("description"), str) else item.get("description") or "",
                "visible": item.get("visible", True),
                "link": item.get("link") or "",
                "linkLabel": item.get("linkLabel") or "",
            }
            for item in args["projects"] if isinstance(item, dict)
        ]

    if args.get("skills") is not None:
        existing["skillContent"] = _ta_to_list(args["skills"]) if args["skills"] else ""

    if args.get("self_evaluation") is not None:
        existing["selfEvaluationContent"] = _ta_to_para(args["self_evaluation"]) if args["self_evaluation"] else ""

    row.data_json = json.dumps(existing, ensure_ascii=False)
    db.commit()
    db.refresh(row)
    return {
        "status": "completed",
        "tool": "update_resume_data",
        "summary": f"简历《{row.title}》已更新，请点击链接进入编辑器查看。",
        "resume_id": row.id,
        "editor_url": f"/student/resumes/{row.id}",
        "open_resume_editor": True,
    }


def _safe_pdf_filename(name: str) -> str:
    base = Path(name.strip()).name or "优化简历"
    if base.lower().endswith(".pdf"):
        base = base[:-4]
    base = "".join(ch for ch in base if ch not in '\\/:*?"<>|').strip() or "优化简历"
    return f"{base[:60]}.pdf"


def _attachment_download_url(stored_path: Path | str) -> str:
    s = str(stored_path).replace("\\", "/")
    marker = "agent_uploads/"
    idx = s.find(marker)
    return "/data/" + (s[idx:] if idx >= 0 else Path(s).name)


def _resolve_student_photo(db: Session, identity: AuthIdentity) -> Optional[str]:
    """解析学生头像的本地文件路径，用于简历照片；找不到/非图片则返回 None。"""
    student = db.get(StudentUser, identity.user_id)
    avatar_url = getattr(student, "avatar_url", None) if student else None
    if not avatar_url:
        return None
    name = Path(str(avatar_url)).name
    if Path(name).suffix.lower() not in (".png", ".jpg", ".jpeg", ".webp"):
        return None
    for base in ("/app/data/avatars", "data/avatars", "./data/avatars"):
        candidate = Path(base) / name
        if candidate.exists():
            return str(candidate)
    return None


def _export_resume_pdf_tool(
    db: Session,
    identity: AuthIdentity,
    session: StudentAgentSession,
    assistant_message: StudentAgentMessage,
    args: dict[str, Any],
) -> dict[str, Any]:
    markdown = str(args.get("markdown") or args.get("content") or "").strip()
    if not markdown:
        return {"status": "failed", "tool": "export_resume_pdf", "summary": "导出失败：未提供简历正文（markdown）。"}

    filename = _safe_pdf_filename(str(args.get("filename") or "优化简历"))
    settings = get_settings()
    storage_dir = Path(settings.agent_upload_storage_dir) / str(identity.tenant_id) / str(identity.user_id)
    storage_dir.mkdir(parents=True, exist_ok=True)
    stored_path = storage_dir / f"{uuid.uuid4().hex}.pdf"
    photo_path = _resolve_student_photo(db, identity)

    try:
        _render_resume_pdf(markdown, stored_path, title=Path(filename).stem, photo_path=photo_path)
    except Exception as exc:  # noqa: BLE001
        return {"status": "failed", "tool": "export_resume_pdf", "summary": f"PDF 生成失败：{str(exc)[:160]}"}

    size = stored_path.stat().st_size if stored_path.exists() else 0
    row = StudentAgentAttachment(
        tenant_id=identity.tenant_id,
        student_id=identity.user_id,
        session_id=session.id,
        message_id=assistant_message.id,
        original_name=filename,
        stored_path=str(stored_path),
        content_type="application/pdf",
        file_ext="pdf",
        file_size=size,
        extracted_text=markdown[:8000],
        status="ready",
    )
    db.add(row)
    db.commit()
    db.refresh(row)

    download_url = _attachment_download_url(stored_path)
    return {
        "status": "completed",
        "tool": "export_resume_pdf",
        "summary": f"已生成简历 PDF：{filename}",
        "filename": filename,
        "download_url": download_url,
        "attachment_id": row.id,
    }


def _pdf_inline(text: str) -> str:
    """Escape for reportlab Paragraph and convert minimal Markdown inline marks."""
    import re

    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
    safe = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", safe)
    return safe


# 候选 CJK 字体（按优先级）。Docker 镜像装了 fonts-noto-cjk；macOS 本地开发用系统字体。
# .ttc 需要 subfontIndex。优先嵌入真实字形，保证任意查看器都能正确渲染中文。
_CJK_FONT_CANDIDATES: tuple[tuple[str, int], ...] = (
    ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
    ("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf", 0),
    ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", 0),
    ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0),
    ("/System/Library/Fonts/PingFang.ttc", 0),
    ("/System/Library/Fonts/STHeiti Light.ttc", 0),
    ("/System/Library/Fonts/Supplemental/Songti.ttc", 0),
)


def _register_cjk_font() -> str:
    """Embed a real CJK font when available (renders everywhere); fall back to the
    non-embedded STSong-Light CID font, and finally to Helvetica."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "ResumeCJK"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name
    for path, subfont_index in _CJK_FONT_CANDIDATES:
        if not Path(path).exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, path, subfontIndex=subfont_index))
            return font_name
        except Exception:
            continue
    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


_ACCENT = "#34507A"  # 板块标题左侧竖条 / 图标 的主题色


def _contact_icon_kind(text: str) -> str:
    """根据联系方式文本推断图标类型。"""
    import re
    t = (text or "").strip()
    low = t.lower()
    if "@" in t:
        return "mail"
    if low.startswith("http") or "www." in low or "://" in low:
        return "globe"
    if re.match(r"^\d{4}[-/.]\d{1,2}", t):
        return "calendar"
    if any(k in t for k in ("离职", "在职", "求职", "在校", "应届", "实习", "全职", "兼职")):
        return "briefcase"
    if re.fullmatch(r"[\d\-\s+()]{7,}", t):
        return "phone"
    return "pin"


def _resume_icon(kind: str, color: str):
    """用矢量图形画一个 12x12 的简约线性图标。"""
    from reportlab.graphics.shapes import Circle, Drawing, Ellipse, Line, Polygon, Rect

    d = Drawing(12, 12)
    sw = 0.9

    def rect(x, y, w, h, **kw):
        return Rect(x, y, w, h, strokeColor=color, strokeWidth=sw, fillColor=None, **kw)

    def line(x1, y1, x2, y2):
        return Line(x1, y1, x2, y2, strokeColor=color, strokeWidth=sw)

    def circ(cx, cy, r):
        return Circle(cx, cy, r, strokeColor=color, strokeWidth=sw, fillColor=None)

    if kind == "mail":
        d.add(rect(1, 2.5, 10, 7))
        d.add(line(1, 9.5, 6, 5.8)); d.add(line(11, 9.5, 6, 5.8))
    elif kind == "phone":
        d.add(rect(3.3, 1, 5.4, 10, rx=1.2, ry=1.2))
        d.add(line(5, 2.3, 7, 2.3))
    elif kind == "calendar":
        d.add(rect(1, 1.5, 10, 8.5))
        d.add(line(1, 7.6, 11, 7.6))
        d.add(line(3.6, 9.8, 3.6, 11.4)); d.add(line(8.4, 9.8, 8.4, 11.4))
    elif kind == "briefcase":
        d.add(rect(1, 2.3, 10, 6.6))
        d.add(rect(4.2, 8.6, 3.6, 1.8))
        d.add(line(1, 5.3, 11, 5.3))
    elif kind == "globe":
        d.add(circ(6, 6, 5))
        d.add(line(1, 6, 11, 6))
        d.add(Ellipse(6, 6, 2.3, 5, strokeColor=color, strokeWidth=sw, fillColor=None))
    else:  # pin
        d.add(circ(6, 8, 3.1))
        d.add(Polygon([3.4, 6.6, 8.6, 6.6, 6, 1], strokeColor=color, strokeWidth=sw, fillColor=None))
        d.add(circ(6, 8, 1.1))
    return d


def _render_resume_pdf(
    markdown_text: str, out_path: Path, title: str = "个人简历", photo_path: Optional[str] = None
) -> None:
    """把约定格式的 Markdown 简历渲染成「专业模板」PDF：左上照片 + 姓名 + 带图标的两列联系方式，
    蓝色竖条 + 灰底的板块标题，三栏对齐（标题/角色/日期）的经历条目，要点带项目符号。
    不符合约定的内容会按通用 Markdown 优雅降级，永不报错。
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import HRFlowable, Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = _register_cjk_font()
    accent = colors.HexColor(_ACCENT)
    content_w = A4[0] - 32 * mm  # 左右各 16mm 边距

    name_st = ParagraphStyle("name", fontName=font_name, fontSize=20, leading=25)
    title_st = ParagraphStyle("title", fontName=font_name, fontSize=10.5, leading=15, textColor=colors.HexColor("#666666"))
    contact_st = ParagraphStyle("contact", fontName=font_name, fontSize=9, leading=13, textColor=colors.HexColor("#444444"))
    sec_st = ParagraphStyle("sec", fontName=font_name, fontSize=11.5, leading=15, textColor=colors.HexColor("#1F2937"))
    entry_l = ParagraphStyle("el", fontName=font_name, fontSize=10.5, leading=14)
    entry_m = ParagraphStyle("em", fontName=font_name, fontSize=10, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#444444"))
    entry_r = ParagraphStyle("er", fontName=font_name, fontSize=9.5, leading=14, alignment=TA_RIGHT, textColor=colors.HexColor("#666666"))
    body = ParagraphStyle("body", fontName=font_name, fontSize=9.8, leading=15, spaceAfter=2)
    bullet = ParagraphStyle("bullet", fontName=font_name, fontSize=9.8, leading=15, leftIndent=10, spaceAfter=1)

    no_pad = [
        ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]

    lines = [ln.rstrip() for ln in markdown_text.splitlines()]
    flow: list[Any] = []

    # ── 头部：照片 + 姓名/职位 + 两列带图标联系方式 ──
    idx = 0
    name = None
    for i, ln in enumerate(lines):
        if ln.strip().startswith("# ") and not ln.strip().startswith("## "):
            name = ln.strip()[2:].strip()
            idx = i + 1
            break
    if name:
        extras: list[str] = []
        while idx < len(lines) and len(extras) < 2:
            s = lines[idx].strip()
            if s.startswith("#"):
                break
            if s:
                extras.append(s)
            idx += 1
        job_title = extras[0] if extras else ""
        contacts = [c.strip() for c in extras[1].split("|") if c.strip()] if len(extras) > 1 else []

        # 左：照片 + 姓名/职位
        name_block = [Paragraph(_pdf_inline(name), name_st)]
        if job_title:
            name_block.append(Paragraph(_pdf_inline(job_title), title_st))
        photo_flow = None
        if photo_path:
            try:
                photo_flow = Image(photo_path, width=46, height=58)
            except Exception:
                photo_flow = None
        if photo_flow is not None:
            left_block: Any = Table([[photo_flow, name_block]], colWidths=[54, content_w * 0.42 - 54])
            left_block.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), *no_pad]))
        else:
            left_block = name_block

        # 右：联系方式两列网格（图标 + 文本）
        cell_w = content_w * 0.58 / 2
        if contacts:
            def contact_cell(text: str) -> Any:
                icon = _resume_icon(_contact_icon_kind(text), _ACCENT)
                inner = Table([[icon, Paragraph(_pdf_inline(text), contact_st)]], colWidths=[15, cell_w - 15])
                inner.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), *no_pad,
                                           ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
                return inner
            grid_rows: list[list[Any]] = []
            for k in range(0, len(contacts), 2):
                grid_rows.append([
                    contact_cell(contacts[k]),
                    contact_cell(contacts[k + 1]) if k + 1 < len(contacts) else "",
                ])
            right_block: Any = Table(grid_rows, colWidths=[cell_w, cell_w])
            right_block.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), *no_pad]))
        else:
            right_block = Paragraph("", contact_st)

        header = Table([[left_block, right_block]], colWidths=[content_w * 0.42, content_w * 0.58])
        header.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), *no_pad]))
        flow.append(header)
        flow.append(Spacer(1, 8))
        flow.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#D0D0D0"), spaceAfter=2))
        rest = lines[idx:]
    else:
        rest = lines  # 没有约定头部 → 整体走通用渲染

    def section_bar(text: str) -> Table:
        # 左侧蓝色竖条 + 灰底标题
        t = Table([["", Paragraph(f"<b>{_pdf_inline(text)}</b>", sec_st)]], colWidths=[3.5, content_w - 3.5])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), accent),
            ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#ECEDEF")),
            ("LEFTPADDING", (0, 0), (0, 0), 0), ("RIGHTPADDING", (0, 0), (0, 0), 0),
            ("LEFTPADDING", (1, 0), (1, 0), 9), ("RIGHTPADDING", (1, 0), (1, 0), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return t

    def entry_row(text: str) -> Any:
        parts = [p.strip() for p in text.split("|")]
        if len(parts) == 1:
            return Paragraph(f"<b>{_pdf_inline(parts[0])}</b>", entry_l)
        if len(parts) == 2:
            cells = [[Paragraph(f"<b>{_pdf_inline(parts[0])}</b>", entry_l), Paragraph(_pdf_inline(parts[1]), entry_r)]]
            widths = [content_w * 0.7, content_w * 0.3]
        else:
            cells = [[
                Paragraph(f"<b>{_pdf_inline(parts[0])}</b>", entry_l),
                Paragraph(_pdf_inline(parts[1]), entry_m),
                Paragraph(_pdf_inline(parts[2]), entry_r),
            ]]
            widths = [content_w * 0.52, content_w * 0.26, content_w * 0.22]
        t = Table(cells, colWidths=widths)
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), *no_pad,
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]))
        return t

    for raw in rest:
        s = raw.strip()
        if not s:
            flow.append(Spacer(1, 3))
        elif s in ("---", "***", "___"):
            flow.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#DDDDDD"), spaceBefore=2, spaceAfter=4))
        elif s.startswith("## "):
            flow.append(Spacer(1, 6))
            flow.append(section_bar(s[3:]))
            flow.append(Spacer(1, 3))
        elif s.startswith("### "):
            flow.append(entry_row(s[4:]))
        elif s.startswith("# "):
            flow.append(Paragraph(f"<b>{_pdf_inline(s[2:])}</b>", name_st))
        elif s[:2] in ("- ", "* ") or s.startswith("• "):
            flow.append(Paragraph("• " + _pdf_inline(s[2:].strip()), bullet))
        else:
            flow.append(Paragraph(_pdf_inline(s), body))

    if not flow:
        flow.append(Paragraph(_pdf_inline(title), name_st))

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=16 * mm, rightMargin=16 * mm, topMargin=16 * mm, bottomMargin=14 * mm, title=title,
    )
    doc.build(flow)

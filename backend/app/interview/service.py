from __future__ import annotations

import json
import re
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import HTTPException, status
from fastapi import UploadFile
from sqlalchemy import select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from app.admin.models import ModelConfig
from app.auth.service import AuthIdentity
from app.core.llm_client import chat_completion
from app.interview.exceptions import (
    InterviewError,
    InterviewLLMError,
    InterviewNoPendingQuestionError,
    InterviewNotActiveError,
    InterviewNotFoundError,
    InterviewReportExistsError,
    InterviewReportGenerationError,
)
from app.interview.harness import (
    SCORE_KEYS,
    _filter_evidence_quotes,
    _strict_bool,
    build_fallback_report,
    harness_should_finish_interview,
    run_harnessed_json_generation,
    validate_followup_output,
    validate_report_output,
    validate_start_output,
)
from app.interview.knowledge import get_knowledge_index, reload_knowledge_index
from app.interview.models import InterviewReport, InterviewSession, InterviewTurn
from app.interview.prompts import (
    EXTRACTED_JOB_PROMPT,
    FOLLOWUP_USER_PROMPT,
    INTERVIEW_FOLLOWUP_SUBPROMPT,
    INTERVIEW_REPORT_SCORING_RUBRIC,
    INTERVIEW_REPORT_SUBPROMPT,
    INTERVIEW_START_SUBPROMPT,
    INTERVIEW_STYLE_CONFIG,
    INTERVIEW_SYSTEM_PROMPT,
    INTERVIEW_TYPE_CONFIG,
    QUALITY_FEEDBACK_PROMPT,
    REPORT_USER_PROMPT,
    SCORING_RUBRIC,
    START_USER_PROMPT,
)
from app.interview.schemas import InterviewStartRequest
from app.student.resume_models import StudentResume


SCORE_WEIGHTS = {
    "technical_accuracy": 0.25,
    "project_evidence": 0.20,
    "problem_solving": 0.20,
    "communication": 0.15,
    "job_fit": 0.15,
    "pressure_handling": 0.05,
}


def _json_loads(raw: str | None, default):
    if not raw:
        return default
    try:
        return json.loads(raw)
    except Exception:
        return default


def _json_dumps(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False)


def _extract_json(text: str) -> dict[str, Any] | None:
    text = text.strip()
    # 1. 直接解析
    try:
        return json.loads(text)
    except Exception:
        pass
    # 2. 剥离 markdown 代码块 ```json ... ``` 或 ``` ... ```
    fence = re.search(r"```(?:json)?\s*\n?([\s\S]*?)```", text)
    if fence:
        inner = fence.group(1).strip()
        try:
            return json.loads(inner)
        except Exception:
            pass
    # 3. 正则提取最外层 JSON 对象（贪婪）
    match = re.search(r"\{[\s\S]*\}", text)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            pass
    # 4. 尝试修复截断的 JSON：逐层补全缺失的 } ]
    for stripped in [text, fence.group(1).strip() if fence else ""]:
        if not stripped or not stripped.startswith("{"):
            continue
        depth = 0
        for ch in stripped:
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
        if depth > 0:
            candidate = stripped + "}" * depth
            try:
                return json.loads(candidate)
            except Exception:
                pass
    return None


def _render_template(template: str, values: dict[str, Any]) -> str:
    """单次遍历替换模板变量，避免已注入值中的 {key} 被后续迭代误替换。"""
    def _replacer(match: re.Match) -> str:
        key = match.group(1)
        return str(values[key]) if key in values else match.group(0)
    return re.sub(r"\{(\w+)\}", _replacer, template)


def _serialize_session(session: InterviewSession) -> dict:
    return {
        "id": session.id,
        "target_role": session.target_role,
        "interview_type": session.interview_type,
        "interview_style": session.interview_style,
        "difficulty": session.difficulty,
        "round_limit": session.round_limit,
        "model_config_id": session.model_config_id,
        "status": session.status,
        "company_name": session.company_name,
        "seniority_level": session.seniority_level,
        "job_skills": _json_loads(session.job_skills_json, []),
        "current_stage": session.current_stage or "opening",
        "created_at": session.created_at.isoformat() if session.created_at else None,
        "ended_at": session.ended_at.isoformat() if session.ended_at else None,
    }


def serialize_turn(turn: InterviewTurn) -> dict:
    return {
        "id": turn.id,
        "turn_index": turn.turn_index,
        "question": turn.question,
        "answer": turn.answer,
        "answer_assessment": _json_loads(turn.answer_assessment, None),
        "score": _json_loads(turn.score_json, None),
        "followup_reason": turn.followup_reason,
        "retrieved_chunks": _json_loads(turn.retrieved_chunks_json, []),
        "knowledge_points": _json_loads(turn.knowledge_points_json, []),
        # 阶段 + 检索解释性 + 评分可解释性
        "stage": turn.stage,
        "question_type": turn.question_type,
        "question_reason": turn.question_reason,
        "capability_tags": _json_loads(turn.capability_tags_json, []),
        "score_reasons": _json_loads(turn.score_reasons_json, {}),
        "evidence_quotes": _json_loads(turn.evidence_quotes_json, []),
        "top_sources": _json_loads(turn.top_sources_json, []),
    }


def serialize_report(report: InterviewReport) -> dict:
    return {
        "id": report.id,
        "session_id": report.session_id,
        "overall_score": report.overall_score,
        "dimension_scores": _json_loads(report.dimension_scores_json, {}),
        "strengths": _json_loads(report.strengths_json, []),
        "weaknesses": _json_loads(report.weaknesses_json, []),
        "suggestions": _json_loads(report.suggestions_json, []),
        "next_questions": _json_loads(report.next_questions_json, []),
        "comparison": _json_loads(report.comparison_json, None),
        "report_text": report.report_text,
        # 训练闭环
        "training_plan": _json_loads(report.training_plan_json, []),
        "rewrite_examples": _json_loads(report.rewrite_examples_json, []),
        "next_session_preset": _json_loads(report.next_session_preset_json, {}),
        "created_at": report.created_at.isoformat() if report.created_at else None,
    }


def knowledge_status() -> dict:
    info = get_knowledge_index().status()
    # 不向学生端暴露服务器绝对路径
    info.pop("root", None)
    return info


def reload_knowledge_status() -> dict:
    info = reload_knowledge_index()
    # 不向学生端暴露服务器绝对路径
    info.pop("root", None)
    return info


def _latest_resume_snapshot(db: Session, identity: AuthIdentity) -> str:
    # 优先读取「智能体可读取」（visibility=True）的简历
    resume = db.scalar(
        select(StudentResume)
        .where(
            StudentResume.student_id == identity.user_id,
            StudentResume.tenant_id == identity.tenant_id,
            StudentResume.visibility.is_(True),
        )
        .limit(1)
    )
    # 若没有开启可读取，则回退到最新保存的简历
    if not resume:
        resume = db.scalar(
            select(StudentResume)
            .where(StudentResume.student_id == identity.user_id, StudentResume.tenant_id == identity.tenant_id)
            .order_by(StudentResume.updated_at.desc())
            .limit(1)
        )
    if not resume:
        return "学生暂未在「简历制作」中保存在线简历。面试时需要优先询问项目、技能和求职方向，并降低对简历证据的确信度。"
    return resume.data_json[:8000]


def _resume_source_label(source: str) -> str:
    if source == "upload":
        return "本次上传简历"
    return "智能体可读取的在线简历"


async def extract_uploaded_resume(upload: UploadFile) -> dict[str, Any]:
    original_name = upload.filename or "resume"
    ext = Path(original_name).suffix.lower()
    if ext not in {".pdf", ".docx", ".txt", ".md"}:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="仅支持 PDF、DOCX、TXT、Markdown 简历")
    content = await upload.read()
    if len(content) > 8 * 1024 * 1024:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="简历文件不能超过 8MB")
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)
    try:
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(tmp_path))
            chunks = []
            for index, page in enumerate(reader.pages[:12], start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    chunks.append(f"[PDF 第 {index} 页]\n{page_text}")
            text = "\n\n".join(chunks)
        elif ext == ".docx":
            from docx import Document

            doc = Document(str(tmp_path))
            chunks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables[:6]:
                for row in table.rows[:24]:
                    values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    if any(values):
                        chunks.append(" | ".join(values))
            text = "\n".join(chunks)
        else:
            text = tmp_path.read_text(encoding="utf-8", errors="ignore")
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
    text = text.strip()[:12000]
    if not text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="未能从简历中提取到可读文本")
    return {
        "filename": original_name,
        "chars": len(text),
        "estimated_tokens": max(1, round(len(text) / 1.5)),
        "extracted_text": text,
    }


def _candidate_chat_models(
    db: Session,
    identity: AuthIdentity,
    preferred_model_id: int | None = None,
) -> list[ModelConfig]:
    models = list(db.scalars(
        select(ModelConfig)
        .where(
            ModelConfig.tenant_id == identity.tenant_id,
            ModelConfig.is_deleted.is_(False),
            ModelConfig.status == "active",
            ModelConfig.open_to_student.is_(True),
            ModelConfig.api_key_cipher.is_not(None),
            ModelConfig.capability.in_(("chat", "text", "multimodal")),
        )
        .order_by(ModelConfig.open_to_student.desc(), ModelConfig.capability.asc(), ModelConfig.id.asc())
    ).all())
    if preferred_model_id:
        models.sort(key=lambda item: 0 if item.id == preferred_model_id else 1)
    return models


def _llm_json(
    db: Session,
    user_prompt: str,
    fallback: dict[str, Any],
    *,
    identity: AuthIdentity | None = None,
    temperature: float = 0.35,
    preferred_model_id: int | None = None,
    max_tokens: int = 2500,
) -> tuple[dict[str, Any], dict[str, Any]]:
    models = _candidate_chat_models(db, identity, preferred_model_id)
    if not models:
        return fallback, {"used": False, "model": None, "error": "No student-open chat model with API key"}
    errors: list[str] = []
    for model in models:
        try:
            result = chat_completion(
                model,
                system_prompt=INTERVIEW_SYSTEM_PROMPT,
                variables={},
                memory=[],
                user_message=user_prompt,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            parsed = _extract_json(result["reply"])
            if not parsed:
                errors.append(f"{model.display_name}: invalid JSON")
                continue
            return parsed, {"used": True, "model": model.display_name, "usage": result.get("usage")}
        except Exception as exc:
            errors.append(f"{model.display_name}: {str(exc)[:180]}")
    return fallback, {"used": False, "model": models[0].display_name, "error": " | ".join(errors)[:500]}


def delete_report(db: Session, identity: AuthIdentity, session_id: int) -> None:
    """删除已有报告，允许重新生成。"""
    session = _get_session(db, identity, session_id)
    db.query(InterviewReport).filter(InterviewReport.session_id == session.id).delete()
    db.commit()


def _conversation_history(turns: list[InterviewTurn]) -> str:
    lines: list[str] = []
    for turn in turns:
        lines.append(f"Q{turn.turn_index}: {turn.question}")
        if turn.answer:
            lines.append(f"A{turn.turn_index}: {turn.answer}")
    return "\n".join(lines[-16:])


def _score_to_100(scores: dict[str, Any]) -> dict[str, float]:
    normalized = {}
    for key in SCORE_KEYS:
        try:
            val = float(scores.get(key, 3))
        except Exception:
            val = 3
        normalized[key] = round(max(1, min(5, val)) * 20, 1)
    return normalized


def _weighted_overall(dim_scores: dict[str, Any]) -> float:
    total = 0.0
    for key, weight in SCORE_WEIGHTS.items():
        try:
            value = float(dim_scores.get(key, 0))
        except Exception:
            value = 0
        total += max(0, min(100, value)) * weight
    return round(total, 1)


def _normalize_report_dimensions(raw: Any, fallback: dict[str, float]) -> dict[str, float]:
    if not isinstance(raw, dict):
        return fallback
    normalized: dict[str, float] = {}
    for key in SCORE_KEYS:
        try:
            value = float(raw.get(key))
        except Exception:
            value = fallback.get(key, 60.0)
        normalized[key] = round(max(0, min(100, value)), 1)
    return normalized


def _fallback_followup(answer: str, retrieved: list[dict]) -> dict[str, Any]:
    topic = retrieved[0]["topic"] if retrieved else "项目经历"
    vague = len(answer.strip()) < 80
    question = (
        f"你刚才的回答还偏概括。围绕 {topic}，请补充一个你亲自处理过的实现细节："
        "具体问题是什么、你做了什么、结果如何量化？"
        if vague
        else f"你提到了这些内容，但我还需要验证深度。请结合 {topic} 说明一个异常场景或取舍：你当时为什么这么设计？"
    )
    return {
        "answer_assessment": {
            "summary": "回答信息量偏少，需要继续追问可验证细节。" if vague else "回答有一定内容，但仍需验证技术深度和个人贡献。",
            "is_vague": vague,
            "risk_points": ["缺少量化指标"] if vague else ["技术取舍说明不足"],
            "positive_points": ["愿意给出项目或技术线索"],
        },
        "score": {
            "technical_accuracy": 3,
            "project_evidence": 2 if vague else 3,
            "problem_solving": 3,
            "communication": 3,
            "job_fit": 3,
            "pressure_handling": 3,
        },
        "score_reasons": {
            "technical_accuracy": "回答偏概括，缺少技术细节支撑",
            "project_evidence": "缺少量化指标和具体项目细节" if vague else "有项目线索但取舍说明不足",
            "problem_solving": "需要补充问题拆解和方案比较",
            "communication": "表达有方向但结构可以更清晰",
            "job_fit": "需要更多岗位核心技术匹配的证据",
            "pressure_handling": "抗压表现待验证",
        },
        "followup_strategy": "追问项目证据和技术细节",
        "interviewer_tone": "strict",
        "next_question": question,
        "question_reason": "回答信息量偏少，需要追问可验证的项目细节和技术深度" if vague else "回答有一定内容，但仍需验证技术深度和个人贡献",
        "question_type": "project_deep_dive",
        "capability_tags": ["项目证据", "技术深度"],
        "knowledge_points": [topic],
        "should_end": False,
        "stage": "resume_deep_dive",
    }


# ── 岗位画像 ──────────────────────────────────────────────────────────────────

_JOB_SKILL_KEYWORDS = [
    "Java", "Spring", "Spring Boot", "MySQL", "Redis", "Kafka",
    "Elasticsearch", "JVM", "Docker", "Kubernetes", "Linux",
    "React", "Vue", "TypeScript", "Python", "Django", "FastAPI", "Flask",
    "LLM", "RAG", "Agent", "MCP", "Function Calling", "LangChain", "LangGraph",
    "数据结构", "算法", "系统设计", "分布式", "微服务", "缓存", "消息队列", "数据库事务",
]


def _extract_job_skills(jd_text: str, user_skills: list[str]) -> list[str]:
    """从 JD 中提取技能标签，优先使用用户手动填写的内容。"""
    if user_skills:
        return list(dict.fromkeys(s.strip() for s in user_skills if s.strip()))
    if not jd_text:
        return []
    found: list[str] = []
    jd_lower = jd_text.lower()
    for kw in _JOB_SKILL_KEYWORDS:
        if kw.lower() in jd_lower:
            found.append(kw)
    return found


# ── 面试阶段状态机 ──────────────────────────────────────────────────────────────

STAGE_DEFINITIONS: dict[str, dict[str, str]] = {
    "opening": {"label": "开场", "goal": "确认目标岗位与面试类型，建立氛围"},
    "self_intro": {"label": "自我介绍", "goal": "考察候选人的自我认知和表达结构"},
    "resume_deep_dive": {"label": "简历深挖", "goal": "验证项目真实性、个人贡献和量化结果"},
    "technical_core": {"label": "核心技术", "goal": "考察岗位必备技术深度和原理理解"},
    "scenario": {"label": "场景题", "goal": "考察系统设计、业务理解和问题拆解能力"},
    "pressure": {"label": "压力追问", "goal": "考察抗压能力、证据意识和诚实度"},
    "reverse_question": {"label": "反问环节", "goal": "考察候选人对岗位和公司的思考深度"},
    "wrap_up": {"label": "收束复盘", "goal": "总结表现，给出改进方向"},
    "completed": {"label": "已完成", "goal": "面试结束"},
}

_STAGE_ORDER = ["opening", "self_intro", "resume_deep_dive", "technical_core", "scenario", "pressure", "reverse_question", "wrap_up"]


def _build_stage_plan(interview_type: str, round_limit: int, focus_tags: list[str]) -> list[dict]:
    """根据面试类型和轮次生成阶段计划。"""
    # 基础分配：按轮次均匀分配阶段（不含 wrap_up，最后单独加）
    stages = [s for s in _STAGE_ORDER if s != "wrap_up"]
    # 压力面跳过 self_intro
    if interview_type == "stress":
        stages = [s for s in stages if s != "self_intro"]
    # HR 面跳过 technical_core
    if interview_type == "hr":
        stages = [s for s in stages if s not in ("technical_core", "pressure")]

    plan: list[dict] = []
    usable_rounds = max(1, round_limit - 1)  # 最后一轮留给 wrap_up
    per_stage = max(1, usable_rounds // len(stages))
    round_num = 1
    for i, stage in enumerate(stages):
        if i == len(stages) - 1:
            # 最后一个阶段用完剩余轮次（不含 wrap_up）
            end = usable_rounds
        else:
            end = min(round_num + per_stage - 1, usable_rounds)
        rounds = list(range(round_num, end + 1))
        if rounds:
            plan.append({"stage": stage, "rounds": rounds})
        round_num = end + 1
    # wrap_up 固定在最后一轮
    plan.append({"stage": "wrap_up", "rounds": [round_limit]})
    return plan


def _stage_for_turn(stage_plan: list[dict], turn_index: int) -> str:
    """根据 turn_index 查找当前阶段。"""
    for entry in stage_plan:
        if turn_index in entry.get("rounds", []):
            return entry["stage"]
    return "opening"


def _update_coverage(coverage: dict, stage: str, knowledge_points: list[str], score: dict) -> dict:
    """更新阶段覆盖度统计。"""
    if stage not in coverage:
        coverage[stage] = {"turns": 0, "knowledge_points": [], "avg_score": 0, "scores": []}
    entry = coverage[stage]
    entry["turns"] += 1
    for kp in knowledge_points:
        if kp not in entry["knowledge_points"]:
            entry["knowledge_points"].append(kp)
    if isinstance(score, dict):
        vals = [v for v in score.values() if isinstance(v, (int, float))]
        if vals:
            entry["scores"].append(sum(vals) / len(vals))
            entry["avg_score"] = round(sum(entry["scores"]) / len(entry["scores"]), 1)
    return coverage


# ── 回答质量感知的阶段推进 ────────────────────────────────────────────────────

def _compute_answer_quality(answer: str, score: dict | None, assessment: dict | None = None) -> tuple[float, bool, bool]:
    """计算回答质量指标。

    Returns:
        (quality_score, is_vague, lacks_depth)
        quality_score: 0-10 分，is_vague: 回答是否空泛，lacks_depth: 是否缺少深度
    """
    answer_len = len(answer.strip()) if answer else 0
    # 基于长度的粗粒度评分
    if answer_len < 30:
        base = 2.0
    elif answer_len < 80:
        base = 4.0
    elif answer_len < 200:
        base = 6.0
    else:
        base = 7.0
    is_vague = answer_len < 80
    lacks_depth = answer_len < 150
    # 如果有 score，结合评分提升
    if isinstance(score, dict):
        try:
            vals = [float(v) for v in score.values() if isinstance(v, (int, float))]
            if vals:
                avg = sum(vals) / len(vals)
                base = round((base + avg * 2) / 2, 1)
        except Exception:
            pass
    # 如果 Model 判定回答空泛，强制降分
    if isinstance(assessment, dict) and assessment.get("is_vague"):
        is_vague = True
        base = min(base, 4.0)
    return round(min(10, max(0, base)), 1), is_vague, lacks_depth


def _update_quality_metrics(coverage: dict, stage: str, quality_score: float, is_vague: bool) -> dict:
    """在 coverage 中增加回答质量指标。"""
    if stage not in coverage:
        coverage[stage] = {"turns": 0, "knowledge_points": [], "avg_score": 0, "scores": [],
                           "quality_scores": [], "avg_quality": 0, "vague_count": 0}
    entry = coverage[stage]
    entry.setdefault("quality_scores", [])
    entry.setdefault("avg_quality", 0)
    entry.setdefault("vague_count", 0)
    entry["quality_scores"].append(quality_score)
    entry["avg_quality"] = round(sum(entry["quality_scores"]) / len(entry["quality_scores"]), 1)
    if is_vague:
        entry["vague_count"] += 1
    return coverage


def _advance_stage(
    current_stage: str,
    stage_plan: list[dict],
    turn_index: int,
    round_limit: int,
    coverage: dict,
    quality_score: float,
    is_vague: bool,
) -> str:
    """根据回答质量和阶段覆盖度决定是否推进阶段。

    规则：
    1. 如果当前阶段回答质量高（≥7）且该阶段已覆盖至少 2 轮，提前推进。
    2. 如果当前阶段回答连续空泛 2 次（使用 consecutive_vague_count），保持当前阶段继续追问。
    3. 如果当前阶段平均质量 ≥ 6 且已覆盖 ≥ 3 轮，推进到下一阶段。
    4. 最后一轮必须是 wrap_up。
    5. 其他情况按 stage_plan 走。
    """
    # 最后一轮强制 wrap_up
    if turn_index >= round_limit - 1:
        return "wrap_up"

    # 查找当前阶段在 plan 中的位置
    current_stage_idx = -1
    for i, entry in enumerate(stage_plan):
        if entry["stage"] == current_stage:
            current_stage_idx = i
            break

    if current_stage_idx < 0:
        return current_stage

    stage_coverage = coverage.get(current_stage, {})
    turns_in_stage = stage_coverage.get("turns", 0)
    avg_quality = stage_coverage.get("avg_quality", 5)
    # 使用连续空泛计数，而非累计空泛计数
    consecutive_vague_count = stage_coverage.get("consecutive_vague_count", 0)

    # 连续空泛，保持当前阶段
    if consecutive_vague_count >= 2 and turns_in_stage < 4:
        return current_stage

    # 高质量回答 + 已覆盖 2 轮 → 提前推进
    if quality_score >= 7 and turns_in_stage >= 2:
        next_idx = current_stage_idx + 1
        if next_idx < len(stage_plan):
            return stage_plan[next_idx]["stage"]
        return current_stage

    # 平均质量 ≥ 6 + 已覆盖 3 轮 → 推进
    if avg_quality >= 6 and turns_in_stage >= 3:
        next_idx = current_stage_idx + 1
        if next_idx < len(stage_plan):
            return stage_plan[next_idx]["stage"]
        return current_stage

    # 按 plan 走：如果当前 turn 已超出当前阶段的 rounds 范围，推进
    current_rounds = stage_plan[current_stage_idx].get("rounds", [])
    if current_rounds and turn_index > max(current_rounds):
        next_idx = current_stage_idx + 1
        if next_idx < len(stage_plan):
            return stage_plan[next_idx]["stage"]

    return current_stage


def _should_skip_stage(stage: str, interview_type: str) -> bool:
    """判断某些阶段是否应该跳过。"""
    if stage == "self_intro" and interview_type == "stress":
        return True
    if stage == "technical_core" and interview_type == "hr":
        return True
    if stage == "pressure" and interview_type == "hr":
        return True
    return False


# wrap_up 阶段允许的问题类型
_WRAP_UP_QUESTION_TYPES = {"wrap_up", "self_review", "reflection", "summary", "closing", "reverse_question"}

# 技术深挖关键词（wrap_up 阶段不应出现）
_DEEP_DIVE_INDICATORS = [
    "算法", "数据结构", "系统设计", "手写", "实现一下", "代码实现",
    "时间复杂度", "空间复杂度", "设计模式", "源码", "底层原理",
    "分布式事务", "CAP 定理", "一致性哈希", "高并发", "压测",
    "请实现", "请写一个", "请设计", "请手撕",
]


def _is_valid_wrap_up_question(question: str, question_type: str) -> bool:
    """判断 wrap_up 阶段的问题是否合法。

    合法的 wrap_up 问题必须是：
    1. question_type 是 wrap_up / self_review / reflection / summary / closing / reverse_question
    2. 问题不包含技术深挖、算法题、系统设计题等关键词
    """
    if question_type not in _WRAP_UP_QUESTION_TYPES:
        return False
    q_lower = question.lower()
    for indicator in _DEEP_DIVE_INDICATORS:
        if indicator in q_lower:
            return False
    return True


def _get_effective_focus_points(retrieved: list[dict]) -> list[str]:
    """从检索结果中提取有效的 focus_points。"""
    points = []
    for item in retrieved[:3]:
        topic = item.get("topic", "")
        if topic and topic not in points:
            points.append(topic)
    return points or ["项目经历", "技术深度", "岗位匹配"]


def _extract_job_profile_info(session: InterviewSession) -> dict:
    """从 session 中提取岗位画像信息。"""
    job_skills = _json_loads(session.job_skills_json, [])
    return {
        "title": session.target_role or "未知岗位",
        "company": session.company_name or "未提供",
        "level": session.seniority_level or "未提供",
        "skills": "、".join(job_skills) if job_skills else "未指定",
        "responsibility": session.job_description[:500] if session.job_description else "未提供",
        "requirements": session.job_description[:500] if session.job_description else "未提供",
    }


# ── wrap_up 本地 fallback ────────────────────────────────────────────────────

def _wrap_up_fallback(target_role: str) -> dict[str, Any]:
    """wrap_up 阶段 LLM 不可用时的本地 fallback。"""
    return {
        "answer_assessment": {
            "summary": "候选人的整体回答表现需要综合评估。",
            "is_vague": False,
            "risk_points": ["需要在报告中综合评估"],
            "positive_points": ["完成了完整的面试流程"],
        },
        "score": {
            "technical_accuracy": 3,
            "project_evidence": 3,
            "problem_solving": 3,
            "communication": 3,
            "job_fit": 3,
            "pressure_handling": 3,
        },
        "score_reasons": {
            "technical_accuracy": "最后一轮综合评估",
            "project_evidence": "最后一轮综合评估",
            "problem_solving": "最后一轮综合评估",
            "communication": "最后一轮综合评估",
            "job_fit": "最后一轮综合评估",
            "pressure_handling": "最后一轮综合评估",
        },
        "evidence_quotes": [],
        "followup_strategy": "收束面试",
        "interviewer_tone": "friendly",
        "next_question": f"感谢你参加这次{target_role}的面试。请用 2 分钟总结一下：你认为自己表现最好的是哪个环节？哪个环节还可以做得更好？",
        "question_reason": "作为收束问题，让候选人自我复盘",
        "question_type": "wrap_up",
        "capability_tags": ["自我认知", "复盘能力"],
        "knowledge_points": [],
        "should_end": True,
        "stage": "wrap_up",
    }


# ── 评分可解释性 ──────────────────────────────────────────────────────────────

def _normalize_score_reasons(raw: Any) -> dict[str, str]:
    """补齐缺失维度的评分原因。"""
    if not isinstance(raw, dict):
        raw = {}
    return {key: str(raw.get(key, "本轮未提供足够证据。")) for key in SCORE_KEYS}


# ── 训练闭环 ──────────────────────────────────────────────────────────────────

def _build_fallback_training_plan(weakest_dim: str) -> list[dict]:
    """当 LLM 未返回训练计划时生成 fallback。"""
    dim_label = {
        "technical_accuracy": "技术准确性",
        "project_evidence": "项目证据",
        "problem_solving": "问题拆解",
        "communication": "表达能力",
        "job_fit": "岗位匹配",
        "pressure_handling": "抗压能力",
    }.get(weakest_dim, "核心能力")
    return [
        {
            "day": 1,
            "focus": dim_label,
            "tasks": ["复盘本轮最低分问题", "准备一个具体项目案例", "补充量化指标"],
            "expected_output": "一段 2 分钟结构化回答",
        },
        {
            "day": 2,
            "focus": "综合练习",
            "tasks": ["用 STAR 结构重写 3 个常见回答", "准备 2 个技术细节追问的应对"],
            "expected_output": "3 个可直接使用的面试回答模板",
        },
    ]


def start_interview(db: Session, identity: AuthIdentity, payload: InterviewStartRequest) -> dict:
    # ── 目标岗位强制必填 ──
    target_role = (payload.target_role or "").strip()
    if not target_role:
        raise InterviewError(status_code=400, detail="请填写目标岗位")

    # ── 岗位 JD 强制必填 ──
    job_description = (payload.job_description or "").strip()
    if not job_description:
        raise InterviewError(status_code=400, detail="请填写岗位 JD")

    if payload.resume_source == "upload":
        resume_snapshot = (payload.uploaded_resume_text or "").strip()[:12000]
        if not resume_snapshot:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="请选择并上传一份可读取的简历")
    else:
        resume_snapshot = _latest_resume_snapshot(db, identity)
    index = get_knowledge_index()
    retrieved = index.search(
        f"{target_role} {job_description} 面试 项目 技术基础",
        target_role=target_role,
        limit=6,
    )
    type_cfg = INTERVIEW_TYPE_CONFIG.get(payload.interview_type, INTERVIEW_TYPE_CONFIG["technical"])
    style_cfg = INTERVIEW_STYLE_CONFIG.get(payload.interview_style, INTERVIEW_STYLE_CONFIG["strict"])
    resume_source_label = _resume_source_label(payload.resume_source)

    # ── 岗位画像 ──
    job_skills = _extract_job_skills(job_description, list(payload.job_skills))
    company_name = (payload.company_name or "").strip() or None
    seniority_level = (payload.seniority_level or "").strip() or None
    job_profile_parts = [f"岗位：{target_role}"]
    if company_name:
        job_profile_parts.append(f"公司：{company_name}")
    if seniority_level:
        job_profile_parts.append(f"级别：{seniority_level}")
    if job_skills:
        job_profile_parts.append(f"核心技能：{'、'.join(job_skills)}")
    job_profile_summary = "，".join(job_profile_parts)

    # ── 阶段计划 ──
    stage_plan = _build_stage_plan(payload.interview_type, payload.round_limit, list(payload.focus_tags))
    current_stage = "opening"

    fallback_start = {
        "resume_brief": f"已读取{resume_source_label}，将围绕岗位匹配度、项目证据和关键能力进行验证。",
        "focus_points": ["项目真实性与个人职责", "目标岗位核心技术匹配", "量化结果和复盘能力"],
        "first_question": (
            f"{type_cfg['opening']} 当前风格是「{style_cfg['label']}」。"
            f"我已经先读取了{resume_source_label}。请选一个最能证明你适合「{target_role}」的项目，"
            "按背景、你的职责、关键方案、量化结果说清楚。"
        ),
        "knowledge_points": [item["topic"] for item in retrieved[:3]] or ["项目证据", "岗位匹配"],
        "question_reason": f"作为开场问题，要求候选人围绕目标岗位「{target_role}」展示最有说服力的项目经历",
        "question_type": "resume_deep_dive",
        "capability_tags": ["项目证据", "岗位匹配"],
    }

    # Prompt 注入岗位画像信息
    profile_injection = f"\n【岗位画像】{job_profile_summary}" if job_skills or company_name or seniority_level else ""
    effective_focus = _get_effective_focus_points(retrieved)
    start_prompt = _render_template(
        START_USER_PROMPT,
        {
            "target_role": target_role,
            "job_description": job_description + profile_injection,
            "interview_type": type_cfg["label"],
            "interview_type_rule": type_cfg["focus"],
            "interview_style": style_cfg["label"],
            "interview_style_rule": style_cfg["rule"],
            "focus_tags": "、".join(payload.focus_tags[:8]) or "、".join(effective_focus),
            "custom_instruction": payload.custom_instruction or "无",
            "resume_summary": resume_snapshot,
            "retrieved_context": json.dumps(retrieved, ensure_ascii=False),
        },
    )
    start_parsed, start_llm_meta = run_harnessed_json_generation(
        db,
        task_name="start_interview",
        system_prompt=INTERVIEW_SYSTEM_PROMPT + "\n\n" + INTERVIEW_START_SUBPROMPT,
        user_prompt=start_prompt,
        fallback=fallback_start,
        validator=validate_start_output,
        identity=identity,
        preferred_model_id=payload.model_id,
        temperature=0.35,
        max_tokens=2500,
        max_retries=2,
    )
    intro = str(start_parsed.get("first_question") or fallback_start["first_question"])
    knowledge_points = start_parsed.get("knowledge_points") if isinstance(start_parsed.get("knowledge_points"), list) else fallback_start["knowledge_points"]
    question_reason = str(start_parsed.get("question_reason") or fallback_start["question_reason"])
    question_type = str(start_parsed.get("question_type") or fallback_start["question_type"])
    capability_tags = start_parsed.get("capability_tags") if isinstance(start_parsed.get("capability_tags"), list) else fallback_start["capability_tags"]

    # 构建 top_sources（只保留 top 3）
    top_sources = [
        {"title": item.get("title", ""), "topic": item.get("topic", ""), "source_file": item.get("source_file", ""), "score": item.get("score", 0)}
        for item in retrieved[:3]
    ]

    session = InterviewSession(
        tenant_id=identity.tenant_id,
        student_id=identity.user_id,
        target_role=target_role,
        job_description=job_description,
        interview_type=payload.interview_type,
        interview_style=payload.interview_style,
        difficulty=payload.difficulty,
        round_limit=payload.round_limit,
        model_config_id=payload.model_id,
        resume_snapshot=f"【简历来源】{resume_source_label}\n【面试类型】{type_cfg['label']}：{type_cfg['focus']}\n【面试风格】{style_cfg['label']}：{style_cfg['rule']}\n【面试重点】{'、'.join(payload.focus_tags[:8]) or '默认'}\n【用户自定义要求】{payload.custom_instruction or '无'}\n\n【岗位画像】{job_profile_summary}\n\n【简历内容】\n{resume_snapshot}",
        # 岗位画像
        company_name=company_name,
        seniority_level=seniority_level,
        job_skills_json=_json_dumps(job_skills),
        job_profile_json=_json_dumps({"summary": job_profile_summary, "skills": job_skills}),
        # 阶段状态机
        current_stage=current_stage,
        stage_plan_json=_json_dumps(stage_plan),
        coverage_json=_json_dumps({}),
    )
    db.add(session)
    db.flush()
    turn = InterviewTurn(
        session_id=session.id,
        student_id=identity.user_id,
        turn_index=1,
        question=intro,
        answer_assessment=_json_dumps({
            "summary": str(start_parsed.get("resume_brief") or fallback_start["resume_brief"]),
            "positive_points": start_parsed.get("focus_points") if isinstance(start_parsed.get("focus_points"), list) else fallback_start["focus_points"],
            "risk_points": [],
            "llm": start_llm_meta,
            "retrieval": {
                "query": f"{target_role} {job_description} 面试 项目 技术基础"[:500],
                "hit_count": len(retrieved),
                "top_sources": [item.get("source_file") for item in retrieved[:3]],
            },
        }),
        retrieved_chunks_json=_json_dumps(retrieved),
        knowledge_points_json=_json_dumps(knowledge_points),
        # 阶段 + 检索解释性
        stage=current_stage,
        question_type=question_type,
        question_reason=question_reason,
        capability_tags_json=_json_dumps(capability_tags),
        retrieval_query=f"{target_role} {job_description} 面试 项目 技术基础"[:500],
        retrieval_hit_count=len(retrieved),
        top_sources_json=_json_dumps(top_sources),
    )
    db.add(turn)
    db.commit()
    db.refresh(session)
    db.refresh(turn)
    return {"session": _serialize_session(session), "first_turn": serialize_turn(turn), "knowledge_status": knowledge_status()}


def _get_session(db: Session, identity: AuthIdentity, session_id: int) -> InterviewSession:
    session = db.get(InterviewSession, session_id)
    if not session or session.student_id != identity.user_id or session.tenant_id != identity.tenant_id:
        raise InterviewNotFoundError
    return session


def list_interviews(db: Session, identity: AuthIdentity) -> list[dict]:
    sessions = db.scalars(
        select(InterviewSession)
        .where(
            InterviewSession.student_id == identity.user_id,
            InterviewSession.tenant_id == identity.tenant_id,
        )
        .order_by(InterviewSession.created_at.desc())
        .limit(50)
    ).all()
    return [_serialize_session(item) for item in sessions]


def get_interview_detail(db: Session, identity: AuthIdentity, session_id: int) -> dict:
    session = _get_session(db, identity, session_id)
    turns = db.scalars(select(InterviewTurn).where(InterviewTurn.session_id == session.id).order_by(InterviewTurn.turn_index)).all()
    return {"session": _serialize_session(session), "turns": [serialize_turn(item) for item in turns]}


def delete_interview(db: Session, identity: AuthIdentity, session_id: int) -> None:
    session = _get_session(db, identity, session_id)
    # 先删子记录（无外键约束，需手动清理）
    db.query(InterviewTurn).filter(InterviewTurn.session_id == session.id).delete()
    db.query(InterviewReport).filter(InterviewReport.session_id == session.id).delete()
    db.delete(session)
    db.commit()


def submit_turn(
    db: Session,
    identity: AuthIdentity,
    session_id: int,
    answer: str,
    *,
    request_id: str | None = None,
    turn_id: int | None = None,
) -> dict:
    session = _get_session(db, identity, session_id)
    if session.status != "active":
        raise InterviewNotActiveError
    turns = db.scalars(select(InterviewTurn).where(InterviewTurn.session_id == session.id).order_by(InterviewTurn.turn_index)).all()
    
    # ── 幂等保护：先处理 turn_id ──
    target_turn = None
    if turn_id is not None:
        target_turn = db.scalar(
            select(InterviewTurn).where(
                InterviewTurn.id == turn_id,
                InterviewTurn.session_id == session.id,
                InterviewTurn.student_id == identity.user_id,
            )
        )
        if not target_turn:
            raise InterviewError(status_code=404, detail="问题不存在")
    else:
        # 找到当前 pending turn
        target_turn = next((turn for turn in reversed(turns) if not turn.answer), None)
        if not target_turn:
            raise InterviewNoPendingQuestionError

    # ── 幂等保护：同一 request_id 重复提交直接返回已有结果 ──
    if target_turn.answer and target_turn.submit_request_id == request_id:
        # 找到已有的 next_turn
        existing_next = next((t for t in turns if t.turn_index == target_turn.turn_index + 1), None)
        return {
            "current_turn": serialize_turn(target_turn),
            "next_turn": serialize_turn(existing_next) if existing_next else None,
            "is_finished": session.status == "completed",
            "report_id": None,
        }

    # ── 幂等保护：同一 turn 不同 request_id 已回答 → 冲突 ──
    if target_turn.answer and (not request_id or target_turn.submit_request_id != request_id):
        raise InterviewError(status_code=409, detail="该问题已回答，请刷新面试记录")

    # ── 检查 target_turn 是否是当前 pending turn ──
    current = next((turn for turn in reversed(turns) if not turn.answer), None)
    if current and target_turn.id != current.id:
        raise InterviewError(status_code=400, detail="turn_id 与当前待回答问题不匹配，请刷新面试记录")
    # 设置 current 为 target_turn，以便后续代码使用
    current = target_turn

    index = get_knowledge_index()
    retrieval_query = f"{session.target_role} {current.question} {answer}"[:500]
    retrieved = index.search(retrieval_query, target_role=session.target_role, limit=6)
    fallback = _fallback_followup(answer, retrieved)

    # ── 构建 top_sources ──
    top_sources = [
        {"title": item.get("title", ""), "topic": item.get("topic", ""), "source_file": item.get("source_file", ""), "score": item.get("score", 0)}
        for item in retrieved[:3]
    ]

    # ── 岗位画像注入 ──
    job_profile = _extract_job_profile_info(session)
    job_profile_text = _render_template(EXTRACTED_JOB_PROMPT, job_profile)

    # ── 注入当前阶段到 Prompt ──
    current_stage = session.current_stage or "opening"
    stage_def = STAGE_DEFINITIONS.get(current_stage, STAGE_DEFINITIONS["opening"])
    stage_injection = f"\n【当前面试阶段】{stage_def['label']}——{stage_def['goal']}"

    # ── 当前回答质量初判（注入给模型）──
    pre_quality_score, pre_is_vague, pre_lacks_depth = _compute_answer_quality(answer, None, None)
    current_quality_injection = (
        f"\n【当前回答质量初判】\n"
        f"质量分：{pre_quality_score}/10\n"
        f"是否空泛：{'是' if pre_is_vague else '否'}\n"
        f"是否缺少深度：{'是' if pre_lacks_depth else '否'}\n"
        f"{'如果空泛，下一问必须要求候选人补充个人职责、实现细节、量化指标或具体案例。' if pre_is_vague else ''}"
    )

    # ── 构建 context_block ──
    context_parts = [
        f"【目标岗位】{session.target_role}",
        f"【岗位 JD】{(session.job_description or '未提供') + stage_injection}",
        f"【面试类型】{INTERVIEW_TYPE_CONFIG.get(session.interview_type, INTERVIEW_TYPE_CONFIG['technical'])['label']}",
        f"【面试风格】{INTERVIEW_STYLE_CONFIG.get(session.interview_style, INTERVIEW_STYLE_CONFIG['strict'])['label']}——{INTERVIEW_STYLE_CONFIG.get(session.interview_style, INTERVIEW_STYLE_CONFIG['strict'])['rule']}",
        f"【候选人简历摘要】{session.resume_snapshot or '未提供'}",
        job_profile_text,
        f"【上一轮问题】{current.question}",
        f"【候选人上一轮回答】{answer}",
        f"【知识库检索结果】{json.dumps(retrieved, ensure_ascii=False)}",
        f"【已问过的知识点】{', '.join(sum((_json_loads(t.knowledge_points_json, []) for t in turns), []))}",
        current_quality_injection,
    ]
    # 注入上一轮回答的质量反馈（供 Model 参考）
    prev_turns_with_answer = [t for t in turns if t.answer and t.turn_index < current.turn_index]
    if prev_turns_with_answer:
        last_turn = prev_turns_with_answer[-1]
        prev_score_data = _json_loads(last_turn.score_json, None)
        prev_assessment_data = _json_loads(last_turn.answer_assessment, None)
        if prev_score_data and prev_assessment_data:
            prev_quality, prev_vague, prev_lacks = _compute_answer_quality(
                last_turn.answer, prev_score_data, prev_assessment_data
            )
            feedback_text = _render_template(QUALITY_FEEDBACK_PROMPT, {
                "quality_score": prev_quality,
                "is_vague": "是" if prev_vague else "否",
                "lacks_depth": "是" if prev_lacks else "否",
                "feedback": "回答空泛，需要更具体的细节和量化指标。" if prev_vague else "",
            })
            context_parts.append("【上一轮已完成回答质量反馈】\n" + feedback_text)
    context_block = "\n\n".join(context_parts)
    conversation_block = f"【历史问答】\n{_conversation_history(turns)}"

    # ── 选择任务 sub-prompt ──
    task_subprompt = INTERVIEW_FOLLOWUP_SUBPROMPT

    prompt = _render_template(
        FOLLOWUP_USER_PROMPT,
        {
            "task_subprompt": task_subprompt,
            "context_block": context_block,
            "conversation_block": conversation_block,
        },
    )

    # 构建 grounding context（供 Harness 校验 next_question 引用式幻觉）
    grounding_context = {
        "last_answer": answer,
        "resume_snapshot": session.resume_snapshot or "",
        "history_text": _conversation_history(turns),
        "job_description": session.job_description or "",
    }

    parsed, llm_meta = run_harnessed_json_generation(
        db,
        task_name="submit_turn",
        system_prompt=INTERVIEW_SYSTEM_PROMPT,
        user_prompt=prompt,
        fallback=fallback,
        validator=validate_followup_output,
        context=grounding_context,
        identity=identity,
        preferred_model_id=session.model_config_id,
        temperature=0.35,
        max_tokens=2500,
        max_retries=2,
    )
    score = parsed.get("score") if isinstance(parsed.get("score"), dict) else fallback["score"]
    assessment = parsed.get("answer_assessment") if isinstance(parsed.get("answer_assessment"), dict) else fallback["answer_assessment"]
    knowledge_points = parsed.get("knowledge_points") if isinstance(parsed.get("knowledge_points"), list) else fallback["knowledge_points"]

    # ── 评分可解释性 ──
    score_reasons = _normalize_score_reasons(parsed.get("score_reasons"))
    evidence_quotes = _filter_evidence_quotes(parsed.get("evidence_quotes"), answer)

    # ── 计算回答质量指标 ──
    quality_score, is_vague, lacks_depth = _compute_answer_quality(answer, score, assessment)

    current.answer = answer
    current.submit_request_id = request_id
    if isinstance(assessment, dict):
        assessment["llm"] = llm_meta
        assessment["retrieval"] = {
            "query": retrieval_query,
            "hit_count": len(retrieved),
            "top_sources": [item.get("source_file") for item in retrieved[:3]],
        }
    current.answer_assessment = _json_dumps(assessment)
    current.score_json = _json_dumps(score)
    current.followup_reason = str(parsed.get("followup_strategy") or parsed.get("followup_reason") or fallback["followup_strategy"])
    current.retrieved_chunks_json = _json_dumps(retrieved)
    current.knowledge_points_json = _json_dumps(knowledge_points)
    # 评分可解释性
    current.score_reasons_json = _json_dumps(score_reasons)
    current.evidence_quotes_json = _json_dumps(evidence_quotes)
    # 检索解释性
    current.retrieval_query = retrieval_query
    current.retrieval_hit_count = len(retrieved)
    current.top_sources_json = _json_dumps(top_sources)

    # ── 更新阶段覆盖度 + 质量指标（区分累计空泛和连续空泛）──
    coverage = _json_loads(session.coverage_json, {})
    coverage = _update_coverage(coverage, current_stage, knowledge_points, score)
    coverage = _update_quality_metrics(coverage, current_stage, quality_score, is_vague)
    # 更新连续空泛计数
    if current_stage not in coverage:
        coverage[current_stage] = {}
    stage_cov = coverage[current_stage]
    if is_vague:
        stage_cov["consecutive_vague_count"] = stage_cov.get("consecutive_vague_count", 0) + 1
    else:
        stage_cov["consecutive_vague_count"] = 0
    session.coverage_json = _json_dumps(coverage)

    # ── Harness 主导的停止判定 ──
    model_should_end = _strict_bool(parsed.get("should_end"))
    valid_answer_count = sum(1 for t in turns if t.answer and len(t.answer.strip()) >= 20)
    coverage_for_decision = _json_loads(session.coverage_json, {})
    should_finish, finish_reason = harness_should_finish_interview(
        model_should_end=model_should_end,
        current_turn_index=current.turn_index,
        round_limit=session.round_limit,
        coverage=coverage_for_decision,
        current_stage=current_stage,
        valid_answer_count=valid_answer_count,
    )
    # 将 finish_reason 写入 assessment 的 llm 字段供审计
    if isinstance(assessment, dict):
        if "llm" not in assessment:
            assessment["llm"] = {}
        assessment["llm"]["finish_reason"] = finish_reason
    report_id = None
    next_turn = None
    if should_finish:
        session.status = "completed"
        session.current_stage = "completed"
        session.ended_at = datetime.now(timezone.utc)
        report = generate_report(db, identity, session.id, commit=False)
        report_id = report.id
    else:
        # ── 计算下一阶段（回答质量感知，使用连续空泛）──
        stage_plan = _json_loads(session.stage_plan_json, [])
        previous_stage = current_stage
        next_stage = _advance_stage(
            current_stage=current_stage,
            stage_plan=stage_plan,
            turn_index=current.turn_index + 1,
            round_limit=session.round_limit,
            coverage=coverage,
            quality_score=quality_score,
            is_vague=is_vague,
        )
        # 跳过不适用的阶段
        while _should_skip_stage(next_stage, session.interview_type) and next_stage != "wrap_up":
            idx = _STAGE_ORDER.index(next_stage) if next_stage in _STAGE_ORDER else -1
            if idx >= 0 and idx + 1 < len(_STAGE_ORDER):
                next_stage = _STAGE_ORDER[idx + 1]
            else:
                break
        session.current_stage = next_stage

        next_question = str(parsed.get("next_question") or fallback["next_question"])
        next_question_type = str(parsed.get("question_type") or fallback.get("question_type", ""))
        next_question_reason = str(parsed.get("question_reason") or fallback.get("followup_strategy", ""))
        next_capability_tags = parsed.get("capability_tags") if isinstance(parsed.get("capability_tags"), list) else []

        # wrap_up 阶段强制收束：不只是 next_question 为空时 fallback
        if next_stage == "wrap_up":
            if not _is_valid_wrap_up_question(next_question, next_question_type):
                wrap_fallback = _wrap_up_fallback(session.target_role)
                next_question = wrap_fallback["next_question"]
                next_question_type = wrap_fallback.get("question_type", "wrap_up")
                next_question_reason = wrap_fallback.get("question_reason", "")
                next_capability_tags = wrap_fallback.get("capability_tags", [])

        # 幂等保护：创建下一轮 turn 前检查 (session_id, turn_index) 是否已存在
        next_turn_index = current.turn_index + 1
        existing_next = db.scalar(
            select(InterviewTurn).where(
                InterviewTurn.session_id == session.id,
                InterviewTurn.turn_index == next_turn_index,
            )
        )
        if existing_next:
            # 已存在，直接复用（防止并发重复创建）
            next_turn = existing_next
            next_turn.question = next_question
            next_turn.stage = next_stage
            next_turn.question_type = next_question_type
            next_turn.question_reason = next_question_reason
        else:
            next_turn = InterviewTurn(
                session_id=session.id,
                student_id=identity.user_id,
                turn_index=next_turn_index,
                question=next_question,
                retrieved_chunks_json=_json_dumps(retrieved),
                knowledge_points_json=_json_dumps(knowledge_points),
                # 阶段
                stage=next_stage,
                question_type=next_question_type,
                question_reason=next_question_reason,
                capability_tags_json=_json_dumps(next_capability_tags),
            )
            db.add(next_turn)
            try:
                db.flush()
            except IntegrityError:
                db.rollback()
                # 回查已有的 next_turn
                existing_next = db.scalar(
                    select(InterviewTurn).where(
                        InterviewTurn.session_id == session.id,
                        InterviewTurn.turn_index == next_turn_index,
                    )
                )
                if existing_next:
                    next_turn = existing_next
                    next_turn.question = next_question
                    next_turn.stage = next_stage
                    next_turn.question_type = next_question_type
                    next_turn.question_reason = next_question_reason
                else:
                    raise
        # 记录阶段推进信息到 answer_assessment（审计用）
        if isinstance(assessment, dict):
            assessment.setdefault("stage_transition", {})
            assessment["stage_transition"] = {
                "from": previous_stage,
                "to": next_stage,
                "quality_score": quality_score,
                "is_vague": is_vague,
                "lacks_depth": lacks_depth,
                "consecutive_vague_count": coverage.get(current_stage, {}).get("consecutive_vague_count", 0),
            }
            current.answer_assessment = _json_dumps(assessment)
    db.commit()
    db.refresh(current)
    if next_turn:
        db.refresh(next_turn)
    return {
        "current_turn": serialize_turn(current),
        "next_turn": serialize_turn(next_turn) if next_turn else None,
        "is_finished": should_finish,
        "report_id": report_id,
    }


def generate_report(db: Session, identity: AuthIdentity, session_id: int, *, commit: bool = True) -> InterviewReport:
    session = _get_session(db, identity, session_id)
    existing = db.scalar(select(InterviewReport).where(InterviewReport.session_id == session.id).order_by(InterviewReport.id.desc()).limit(1))
    if existing:
        return existing
    turns = db.scalars(select(InterviewTurn).where(InterviewTurn.session_id == session.id).order_by(InterviewTurn.turn_index)).all()
    scores = [_json_loads(turn.score_json, {}) for turn in turns if turn.score_json]
    if scores:
        dim_scores = {
            key: round(sum(float(score.get(key, 3)) for score in scores) / len(scores) * 20, 1)
            for key in SCORE_KEYS
        }
    else:
        dim_scores = {key: 60.0 for key in SCORE_KEYS}
    overall = _weighted_overall(dim_scores)
    fallback = {
        "overall_score": overall,
        "dimension_scores": dim_scores,
        "strengths": ["能够完成基本面试对话", "已有部分技术或项目线索可继续深挖"],
        "weaknesses": ["回答需要更多量化指标", "项目个人贡献和技术取舍还需要讲得更具体"],
        "suggestions": ["用 STAR 结构回答项目题", "每个技术点准备一个真实故障或优化案例", "回答优化类问题时补充前后数据"],
        "next_questions": ["请介绍一个你亲自优化过的接口。", "Redis 缓存和数据库一致性如何保证？", "如果系统 QPS 突增 10 倍，你会怎么排查瓶颈？"],
        "report_text": f"本次面试综合分 {overall}。整体表现可以继续打磨，重点补充项目证据、数据指标和技术取舍。面试官会认可诚实和细节，不会认可空泛的'负责'和'熟悉'。",
        "training_plan": _build_fallback_training_plan(
            min(dim_scores, key=lambda k: dim_scores.get(k, 100)) if dim_scores else "project_evidence"
        ),
        "rewrite_examples": [],
        "next_session_preset": {
            "target_role": session.target_role,
            "interview_type": session.interview_type,
            "interview_style": session.interview_style,
            "focus_tags": [],
        },
    }

    # ── 注入阶段覆盖度 ──
    coverage = _json_loads(session.coverage_json, {})
    coverage_summary = ""
    if coverage:
        coverage_lines = []
        for stage_key, info in coverage.items():
            stage_label = STAGE_DEFINITIONS.get(stage_key, {}).get("label", stage_key)
            avg_q = info.get("avg_quality", 0)
            coverage_lines.append(f"  {stage_label}: {info.get('turns', 0)} 轮, 平均分 {info.get('avg_score', 0)}, 回答质量 {avg_q}/10")
        coverage_summary = "\n【阶段覆盖度】\n" + "\n".join(coverage_lines)

    # ── 构建报告 prompt（新模板结构）──
    context_parts = [
        f"【目标岗位】{session.target_role}",
        f"【岗位 JD】{session.job_description or '未提供'}",
        f"【简历快照】{(session.resume_snapshot or '未提供')[:12000]}",
        f"【每轮过程评分，仅作参考】{json.dumps(scores, ensure_ascii=False) + coverage_summary}",
    ]
    prompt = _render_template(
        REPORT_USER_PROMPT,
        {
            "task_subprompt": INTERVIEW_REPORT_SUBPROMPT,
            "scoring_rubric_block": f"【评分 Rubric】\n{INTERVIEW_REPORT_SCORING_RUBRIC}",
            "context_block": "\n\n".join(context_parts),
            "conversation_block": f"【面试记录】\n{_conversation_history(turns)}",
        },
    )
    parsed, llm_meta = run_harnessed_json_generation(
        db,
        task_name="generate_report",
        system_prompt=INTERVIEW_SYSTEM_PROMPT,
        user_prompt=prompt,
        fallback=fallback,
        validator=validate_report_output,
        identity=identity,
        preferred_model_id=session.model_config_id,
        temperature=0.2,
        max_tokens=4200,
        max_retries=3,
    )
    final_dim_scores = _normalize_report_dimensions(parsed.get("dimension_scores"), dim_scores)
    try:
        model_overall = float(parsed.get("overall_score"))
    except Exception:
        model_overall = _weighted_overall(final_dim_scores)
    final_overall = round(max(0, min(100, model_overall)), 1)
    weighted_overall = _weighted_overall(final_dim_scores)
    if abs(final_overall - weighted_overall) > 8:
        final_overall = weighted_overall
    comparison = _build_report_comparison(db, identity, session, final_overall, final_dim_scores)
    if comparison is not None:
        comparison["scoring"] = {
            "mode": "llm_rubric" if llm_meta.get("used") else "local_fallback",
            "model": llm_meta.get("model"),
            "usage": llm_meta.get("usage"),
            "rubric": "CareerForge technical/behavioral interview rubric",
        }
    report_text = str(parsed.get("report_text") or fallback["report_text"])
    if not llm_meta.get("used"):
        report_text += "\n\n本次模型评分服务暂时不可用，系统已按同一套评分 Rubric 做本地兜底；建议模型服务恢复后重新生成报告。"
    report = InterviewReport(
        session_id=session.id,
        student_id=identity.user_id,
        overall_score=final_overall,
        dimension_scores_json=_json_dumps(final_dim_scores),
        strengths_json=_json_dumps(parsed.get("strengths") or fallback["strengths"]),
        weaknesses_json=_json_dumps(parsed.get("weaknesses") or fallback["weaknesses"]),
        suggestions_json=_json_dumps(parsed.get("suggestions") or fallback["suggestions"]),
        next_questions_json=_json_dumps(parsed.get("next_questions") or fallback["next_questions"]),
        comparison_json=_json_dumps(comparison),
        report_text=report_text,
        # 训练闭环
        training_plan_json=_json_dumps(parsed.get("training_plan") or fallback["training_plan"]),
        rewrite_examples_json=_json_dumps(parsed.get("rewrite_examples") or fallback["rewrite_examples"]),
        next_session_preset_json=_json_dumps(parsed.get("next_session_preset") or fallback["next_session_preset"]),
    )
    db.add(report)
    session.status = "completed"
    session.ended_at = session.ended_at or datetime.now(timezone.utc)
    if commit:
        db.commit()
        db.refresh(report)
    else:
        db.flush()
    return report


def _build_report_comparison(
    db: Session,
    identity: AuthIdentity,
    session: InterviewSession,
    overall: float,
    dim_scores: dict[str, float],
) -> dict[str, Any] | None:
    previous = db.scalar(
        select(InterviewReport)
        .join(InterviewSession, InterviewReport.session_id == InterviewSession.id)
        .where(
            InterviewReport.student_id == identity.user_id,
            InterviewSession.tenant_id == identity.tenant_id,
            InterviewReport.session_id != session.id,
            InterviewSession.target_role == session.target_role,
        )
        .order_by(InterviewReport.created_at.desc())
        .limit(1)
    )
    if not previous:
        return {
            "has_previous": False,
            "message": "这是该岗位的首次面试记录，后续报告会自动和上一次对比。",
        }
    prev_dims = _json_loads(previous.dimension_scores_json, {})
    prev_overall = float(previous.overall_score or 0)
    delta = round(overall - prev_overall, 1)
    dim_delta = {
        key: round(float(dim_scores.get(key, 0)) - float(prev_dims.get(key, 0)), 1)
        for key in SCORE_KEYS
    }
    if delta >= 5:
        message = f"比上一次提升了 {delta} 分，表现明显更稳。继续保持，别给面试官挑刺的机会。"
    elif delta <= -5:
        message = f"比上一次下降了 {abs(delta)} 分，主要需要回到项目细节和量化结果上补强。"
    else:
        message = f"和上一次基本持平（{delta:+.1f} 分），下一轮建议集中突破最低分维度。"
    return {
        "has_previous": True,
        "previous_report_id": previous.id,
        "previous_overall_score": prev_overall,
        "current_overall_score": overall,
        "overall_delta": delta,
        "dimension_delta": dim_delta,
        "message": message,
    }


def get_report(db: Session, identity: AuthIdentity, session_id: int) -> dict:
    """Get report for a session."""
    session = _get_session(db, identity, session_id)
    report = db.scalar(
        select(InterviewReport)
        .where(InterviewReport.session_id == session.id)
        .order_by(InterviewReport.id.desc())
        .limit(1)
    )
    if not report:
        raise InterviewError(status_code=404, detail="报告不存在")
    return serialize_report(report)


def export_interview_report(db: Session, identity: AuthIdentity, session_id: int) -> dict:
    """Export full interview report as JSON."""
    session = _get_session(db, identity, session_id)
    turns = db.scalars(
        select(InterviewTurn)
        .where(InterviewTurn.session_id == session.id)
        .order_by(InterviewTurn.turn_index)
    ).all()
    report = db.scalar(
        select(InterviewReport)
        .where(InterviewReport.session_id == session.id)
        .order_by(InterviewReport.id.desc())
        .limit(1)
    )
    return {
        "session": _serialize_session(session),
        "turns": [serialize_turn(t) for t in turns],
        "report": serialize_report(report) if report else None,
    }
